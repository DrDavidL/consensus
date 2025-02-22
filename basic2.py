
#########################################
# Import Libraries and Setup
#########################################
import asyncio
import json
import re
import tempfile
import time
from datetime import datetime, timedelta
from io import BytesIO
import logging

import aiohttp
import requests
import streamlit as st
import markdown2
import xml.etree.ElementTree as ET
from docx import Document
from typing import List, Tuple, Dict

# External libraries for AI and vector store
from openai import OpenAI
import logging
import chromadb
from chromadb.config import Settings

# (Other libraries such as anthropic, exa_py, tavily remain if needed)
import anthropic
from openai import OpenAI
from exa_py import Exa
from tavily import TavilyClient
import chromadb
from chromadb.config import Settings
import requests
from typing import List, Dict

from bs4 import BeautifulSoup

#########################################
# Global Variables and Logging Configuration
#########################################
role_emojis = {
    "user": "👤",
    "assistant": "🤖",
}

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)



#########################################
# Import Prompts for AI Guidance and Search
#########################################
from prompts import (
    system_prompt_expert_questions,
    expert1_system_prompt,
    expert2_system_prompt,
    expert3_system_prompt,
    optimize_search_terms_system_prompt,
    optimize_pubmed_search_terms_system_prompt,
    cutting_edge_pubmed_prompt,
    prepare_rag_query,
    rag_prompt2,
    choose_domain,
    medical_domains,
    prelim_followup_prompt,
    tavily_domains,
)

#########################################
# Streamlit App Configuration and API Keys Setup
#########################################
st.set_page_config(
    page_title="Helpful AI",
    layout="wide",
    page_icon=":stethoscope:",
    initial_sidebar_state="collapsed",
)
api_key = st.secrets["OPENAI_API_KEY"]
api_key_anthropic = st.secrets["ANTHROPIC_API_KEY"]
exa = Exa(st.secrets["EXA_API_KEY"])

#########################################
# Session State Initialization
#########################################
if "snippets" not in st.session_state:
    st.session_state["snippets"] = []
if "urls" not in st.session_state:
    st.session_state["urls"] = []
if "rag_response" not in st.session_state:
    st.session_state["rag_response"] = ""
if "source_chunks" not in st.session_state:
    st.session_state.source_chunks = ""
if "experts" not in st.session_state:
    st.session_state.experts = []
if "messages1" not in st.session_state:
    st.session_state.messages1 = []
if "messages2" not in st.session_state:
    st.session_state.messages2 = []
if "messages3" not in st.session_state:
    st.session_state.messages3 = []
if "followup_messages" not in st.session_state:
    st.session_state.followup_messages = []
if "final_thread" not in st.session_state:
    st.session_state.final_thread = []
if "expert_number" not in st.session_state:
    st.session_state.expert_number = 0
if "expert_answers" not in st.session_state:
    st.session_state.expert_answers = []
if "original_question" not in st.session_state:
    st.session_state.original_question = ""
if "chosen_domain" not in st.session_state:
    st.session_state.chosen_domain = ""
if "articles" not in st.session_state:
    st.session_state.articles = []
if "pubmed_search_terms" not in st.session_state:
    st.session_state.pubmed_search_terms = ""
if "full_initial_response" not in st.session_state:
    st.session_state.full_initial_response = ""
if "initial_response_thread" not in st.session_state:
    st.session_state.initial_response_thread = []
if "citations" not in st.session_state:
    st.session_state.citations = []
if "thread_with_tavily_context" not in st.session_state:
    st.session_state.thread_with_tavily_context = []
if "tavily_followup_response" not in st.session_state:
    st.session_state.tavily_followup_response = ""
if "tavily_initial_response" not in st.session_state:
    st.session_state.tavily_initial_response = ""
if "tavily_urls" not in st.session_state:
    st.session_state.tavily_urls = ""



class OpenAIEmbeddingFunction:
    """
    A wrapper for OpenAI's embedding API using the new client interface.
    This class implements the __call__(self, input) signature required by Chroma.
    """
    def __init__(self, model: str, api_key: str):
        self.model = model
        self.client = OpenAI(api_key=api_key)

    def __call__(self, input):
        """
        Returns embedding(s) for the given input.
        If input is a list, returns a list of embeddings; otherwise returns a single embedding.
        """
        try:
            response = self.client.embeddings.create(
                input=input,
                model=self.model
            )
            # response.data is a list of objects with an 'embedding' attribute.
            if isinstance(input, list):
                return [item.embedding for item in response.data]
            else:
                return response.data[0].embedding
        except Exception as e:
            logger.error(f"Embedding error: {e}")
            return []


#########################################
# Updated MyApp class using new embedding function
#########################################

def is_relevant(text: str) -> bool:
    """
    Determines if a paragraph is relevant based on medical and technical content.
    Ignores metadata-heavy or generic sections.
    """
    irrelevant_patterns = [
        r"^The site is secure",
        r"^An official website of the",
        r"^PubMed Disclaimer",
        r"^Keywords:",
        r"^References",
        r"^Funding:",
        r"^Conflict of Interest",
        r"^National Library of Medicine"
    ]
    
    for pattern in irrelevant_patterns:
        if re.match(pattern, text):
            return False
    
    # Ensure the paragraph has medical/scientific keywords
    relevant_keywords = ["infection", "treatment", "risk", "disease", "pathogenesis", "mortality", "morbidity"]
    return any(word in text.lower() for word in relevant_keywords)


def clean_html(html_content: str) -> str:
    """
    Cleans HTML content by stripping tags, script/style elements, and extra whitespace,
    then segments the plain text while ensuring relevance.
    """
    soup = BeautifulSoup(html_content, "lxml")  # Faster parser

    # Remove script and style elements
    for element in soup(["script", "style"]):
        element.extract()

    # Extract paragraphs while preserving structure
    paragraphs = [p.get_text(separator=" ").strip() for p in soup.find_all("p") if p.get_text(strip=True)]
    
    # If no paragraphs, fall back to whole text
    if not paragraphs:
        text = soup.get_text(separator=" ")
        text = re.sub(r"\s+", " ", text).strip()
        return segment_text([process_chunk(text)], min_size=3000, max_size=5000, separator="\n---\n")

    # Process and segment relevant paragraphs
    processed_paragraphs = [process_chunk(p) for p in paragraphs if is_relevant(p)]
    return segment_text(processed_paragraphs, min_size=3000, max_size=5000, separator="\n---\n")

def segment_text(paragraphs: List[str], min_size=3000, max_size=5000, separator="\n---\n") -> str:
    """
    Segments text into chunks that:
    - Start at the beginning of a topic (not mid-sentence)
    - End at a paragraph boundary
    - Maintain logical groupings of content
    """
    segments = []
    current_chunk = ""

    for paragraph in paragraphs:
        if len(current_chunk) + len(paragraph) < max_size:
            current_chunk += paragraph + "\n\n"
        else:
            if len(current_chunk) >= min_size:
                segments.append(current_chunk.strip())
                current_chunk = paragraph + "\n\n"
            else:
                current_chunk += paragraph + "\n\n"  # If too small, keep adding

    if current_chunk:
        segments.append(current_chunk.strip())  # Append the last chunk

    return separator.join(segments)


def process_chunk(chunk: str) -> str:
    """
    Cleans and compresses a chunk of text by:
    - Removing irrelevant sections (disclaimers, legal notes, website security text)
    - Summarizing verbose sections while maintaining readability
    - Filtering out noise like 'FIG 1' or 'Keywords'
    """
    patterns_to_remove = [
        r"An official website of the United States government.*?",  # Government notices
        r"The site is secure.*?",  # Website security text
        r"PubMed Disclaimer.*?",  # PubMed metadata
        r"Connect with NLM.*?",  # National Library of Medicine metadata
        r"Keywords:.*?",  # Keyword lists
        r"FIG \d+.*?",  # Figure references
        r"MeSH PMC Bookshelf Disclaimer.*?",  # Miscellaneous disclaimers
        r"\bReferences\b.*",  # Reference sections
        r"\bConflicts of Interest\b.*",  # Conflict of interest statements
    ]
    
    # Remove irrelevant sections
    for pattern in patterns_to_remove:
        chunk = re.sub(pattern, "", chunk, flags=re.DOTALL).strip()

    # Summarize core content while keeping medical details
    chunk = compress_text(chunk)
    
    return chunk


def compress_text(text: str, max_length: int = 1000) -> str:
    """
    Compresses text while maintaining readability and core medical content.
    Uses sentence compression and key information extraction.
    """
    sentences = re.split(r'(?<=[.!?])\s+', text)  # Split at sentence boundaries
    summary = []
    word_count = 0
    
    for sentence in sentences:
        if len(sentence.split()) < 5:  # Skip very short, non-informative sentences
            continue
        
        if word_count + len(sentence.split()) > max_length:  # Stop when reaching max length
            break
        
        summary.append(sentence)
        word_count += len(sentence.split())
    
    return " ".join(summary)

class MyApp:
    """
    An app for storing embeddings in a Chroma vectorstore.
    It fetches website content, splits it into chunks, embeds them using the provided OpenAIEmbeddingFunction,
    and saves/retrieves these chunks for semantic search.
    """
    def __init__(self, config: dict):
        self.config = config
        self.collection_name = config["vectordb"]["config"]["collection_name"]
        self.db_dir = config["vectordb"]["config"]["dir"]
        self.embedder_model = config["embedder"]["config"]["model"]
        self.embedder_api_key = config["embedder"]["config"]["api_key"]
        self.chunk_size = config["chunker"]["chunk_size"]
        self.chunk_overlap = config["chunker"]["chunk_overlap"]

        # Create an instance of our embedding function that meets Chroma's requirements.
        self.embedding_fn = OpenAIEmbeddingFunction(self.embedder_model, self.embedder_api_key)

        # Initialize the persistent Chroma client.
        self.client = chromadb.PersistentClient(
            path=self.db_dir,
            settings=Settings(allow_reset=True, anonymized_telemetry=False)
        )
        # Use get_or_create_collection to ensure the collection exists.
        self.collection = self.client.get_or_create_collection(
            name=self.collection_name,
            embedding_function=self.embedding_fn
        )
    

    def chunk_text(self, text: str, separator="\n---\n") -> List[str]:
        """
        Splits text into chunks based on a predefined separator instead of a fixed-size sliding window.
        """
        return text.split(separator)

    def add(self, url: str, data_type: str = "web_page") -> None:
        """
        Fetches content from the URL, cleans the HTML to extract the plain text,
        chunks the text, embeds each chunk, and adds it to the Chroma collection.
        """
        try:
            response = requests.get(url, headers = {"User-Agent": "Mozilla/5.0"}, timeout=10)
            response.raise_for_status()
            raw_content = response.text
            cleaned_content = clean_html(raw_content)
        except Exception as e:
            logger.error(f"Error fetching {url}: {e}")
            return

        chunks = self.chunk_text(cleaned_content)
        for i, chunk in enumerate(chunks):
            embedding = self.embedding_fn(chunk)
            doc_id = f"{url}_{i}"
            self.collection.add(
                documents=[chunk],
                embeddings=[embedding],
                metadatas=[{"url": url, "chunk_index": i, "data_type": data_type}],
                ids=[doc_id],
            )


  
    def search(self, query: str, num_documents: int = 20, where: dict = None) -> List[Dict]:
        """Perform a semantic search using the query and return formatted citations."""
        query_embedding = self.embedding_fn(query)
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=num_documents,
            include=["metadatas", "documents", "distances"],  # 'ids' removed as per new API requirements
            where=where,
        )
        st.write("Raw query results:", results)
        
        citations = []
        
        # Iterate over returned documents
        for i in range(len(results["documents"][0])):
            distance = results["distances"][0][i]
            
            # Convert ChromaDB's cosine distance to a proper similarity score
            similarity_score = 1 - (distance / 2)  # Ensures range 0 (worst) to 1 (best)
            
            citation = {
                "context": results["documents"][0][i],
                "metadata": results["metadatas"][0][i],
                "score": similarity_score
            }
            citations.append(citation)
        
        return citations



    def reset(self) -> None:
        """Reset the vectorstore by deleting and re-creating the collection."""
        self.client.delete_collection(self.collection_name)
        self.collection = self.client.get_or_create_collection(
            name=self.collection_name,
            embedding_function=self.embedding_fn
        )

    @classmethod
    def from_config(cls, config: dict) -> "MyApp":
        return cls(config)


#########################################
# Utility Functions (unchanged)
#########################################
def is_non_informative(context: str) -> bool:
    context = context.strip()
    if not context or len(context) < 5:
        return True
    name_pattern = re.compile(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?$')
    if name_pattern.fullmatch(context):
        return True
    author_et_al_pattern = re.compile(r'^[A-Z][a-z]+\s+et\s+al\.?$', re.IGNORECASE)
    if author_et_al_pattern.fullmatch(context):
        return True
    pubmed_mentions = len(re.findall(r'PubMed:', context, flags=re.IGNORECASE))
    pmc_mentions = len(re.findall(r'PMC', context, flags=re.IGNORECASE))
    if (pubmed_mentions + pmc_mentions) > 5:
        return True
    doi_matches = re.findall(r'\bdoi:\s*\S+', context, re.IGNORECASE)
    pmid_matches = re.findall(r'\bPMID:\s*\d+', context)
    if (len(doi_matches) + len(pmid_matches)) > 5:
        return True
    numbered_ref_pattern = re.compile(r'^\d+\.\s', re.MULTILINE)
    if len(numbered_ref_pattern.findall(context)) > 3:
        return True
    if context.count("et al") > 1 or re.search(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b', context):
        if len(context) > 150:
            return True
    if context.count("Disclosure:") > 1:
        return True

    return False

def filter_citations(citations: list) -> list:
    return [
        {
            "context": citation.get("context", ""),
            "url": citation.get("metadata", {}).get("url", ""),
            "score": citation.get("score", 0),  # Use the top-level score
        }
        for citation in citations
        if not is_non_informative(citation.get("context", ""))
    ]


def extract_and_format_urls(tavily_output):
    results = tavily_output.get("results", [])
    if not results:
        return "No URLs found."
    urls = [result.get("url", "") for result in results if result.get("url", "")]
    unique_urls = sorted(set(urls))
    output_lines = ["List of References:"]
    for idx, url in enumerate(unique_urls, start=1):
        output_lines.append(f"{idx}. {url}")
    return "\n".join(output_lines)

def display_url_list(citations):
    urls = [citation.get("url", "") for citation in citations if citation.get("url", "")]
    unique_urls = sorted(set(urls))
    st.markdown("**List of Source URLs**", unsafe_allow_html=True)
    for url in unique_urls:
        st.markdown(f"- [{url}]({url})", unsafe_allow_html=True)

def display_citations(citations):
    st.markdown("## Sources")
    sorted_citations = sorted(citations, key=lambda c: c.get("score", 0), reverse=True)
    for i, citation in enumerate(sorted_citations, start=1):
        normalized_score = round(citation.get("score", 0) * 100, 2)
        st.markdown(f"### Source {i} (Relevance: {normalized_score}%)", unsafe_allow_html=True)
        url = citation.get("url", "")
        if url:
            st.markdown(f"[Link to Source]({url})", unsafe_allow_html=True)
        context_text = citation.get("context", "")
        st.text(context_text)
        st.markdown("---")

def markdown_to_word(markdown_text):
    doc = Document()
    lines = markdown_text.split("\n")
    for line in lines:
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    return doc

async def extract_abstract_from_xml(xml_data: str, pmid: str) -> str:
    try:
        root = ET.fromstring(xml_data)
        for article in root.findall(".//PubmedArticle"):
            medline_citation = article.find("MedlineCitation")
            if medline_citation:
                pmid_element = medline_citation.find("PMID")
                if pmid_element is not None and pmid_element.text == pmid:
                    abstract_element = medline_citation.find(".//Abstract")
                    if abstract_element is not None:
                        abstract_texts = []
                        for elem in abstract_element.findall("AbstractText"):
                            label = elem.get("Label")
                            text = ET.tostring(elem, encoding="unicode", method="text").strip()
                            if label:
                                abstract_texts.append(f"{label}: {text}")
                            else:
                                abstract_texts.append(text)
                        return " ".join(abstract_texts).strip()
        return "No abstract available"
    except ET.ParseError:
        print(f"Error parsing XML for PMID {pmid}")
        return "Error extracting abstract"

async def fetch_additional_results(session: aiohttp.ClientSession, search_query: str, max_results: int, current_count: int) -> List[str]:
    additional_needed = max_results - current_count
    url = (
        f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?"
        f"db=pubmed&term={search_query}&sort=relevance&retmode=json&retmax={additional_needed}&"
        f"api_key={st.secrets['pubmed_api_key']}"
    )
    try:
        async with session.get(url) as response:
            response.raise_for_status()
            data = await response.json()
            return data["esearchresult"].get("idlist", [])
    except (aiohttp.ClientError, asyncio.TimeoutError) as e:
        print(f"Error fetching additional results: {e}")
        return []

async def fetch_article_details(session: aiohttp.ClientSession, id: str, details_url: str, abstracts_url: str, semaphore: asyncio.Semaphore) -> Tuple[str, Dict, str]:
    async with semaphore:
        try:
            async with session.get(details_url) as details_response:
                details_response.raise_for_status()
                details_data = await details_response.json()
            async with session.get(abstracts_url) as abstracts_response:
                abstracts_response.raise_for_status()
                abstracts_data = await abstracts_response.text()
            return id, details_data, abstracts_data
        except (aiohttp.ClientError, asyncio.TimeoutError) as e:
            print(f"Error fetching article details for ID {id}: {e}")
            return id, {}, ""

async def pubmed_abstracts(search_terms: str, search_type: str = "all", max_results: int = 6, years_back: int = 4, filter_relevance: bool = True, relevance_threshold: float = 0.8) -> List[Dict[str, str]]:
    current_year = datetime.now().year
    start_year = current_year - years_back
    search_query = f"{search_terms}+AND+{start_year}[PDAT]:{current_year}[PDAT]"
    url = (
        f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?"
        f"db=pubmed&term={search_query}&sort=relevance&retmode=json&retmax={max_results}&"
        f"api_key={st.secrets['pubmed_api_key']}"
    )
    async with aiohttp.ClientSession() as session:
        try:
            async with session.get(url) as response:
                response.raise_for_status()
                data = await response.json()
                if "esearchresult" not in data or "count" not in data["esearchresult"]:
                    st.error("Unexpected response format from PubMed API")
                    return []
                ids = data["esearchresult"].get("idlist", [])
                if not ids:
                    st.write("No results found.")
                    return []
            articles = []
            semaphore = asyncio.Semaphore(5)
            tasks = []
            for id in ids:
                details_url = (
                    f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi?"
                    f"db=pubmed&id={id}&retmode=json&api_key={st.secrets['pubmed_api_key']}"
                )
                abstracts_url = (
                    f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?"
                    f"db=pubmed&id={id}&retmode=xml&rettype=abstract&api_key={st.secrets['pubmed_api_key']}"
                )
                tasks.append(fetch_article_details(session, id, details_url, abstracts_url, semaphore))
            results = await asyncio.gather(*tasks)
            filtered_articles = []
            for id, details_data, abstracts_data in results:
                if "result" in details_data and str(id) in details_data["result"]:
                    article = details_data["result"][str(id)]
                    year = article["pubdate"].split(" ")[0]
                    if year.isdigit():
                        abstract = await extract_abstract_from_xml(abstracts_data, id)
                        article_url = f"https://pubmed.ncbi.nlm.nih.gov/{id}"
                        if abstract.strip() and abstract != "No abstract available":
                            filtered_articles.append({
                                "id": id,
                                "title": article["title"],
                                "year": year,
                                "abstract": abstract.strip(),
                                "link": article_url,
                            })
            if filter_relevance:
                articles_prompt = "\n".join(
                    [f"ID: {article['id']} - Title: {article['title']}" for article in filtered_articles]
                )
                messages = [
                    {
                        "role": "system",
                        "content": (
                            "You are an assistant evaluating the relevance of articles to a query. "
                            "For each article provided, return a relevance score between 0 and 1 as a JSON object mapping "
                            "the article's ID to its score."
                        ),
                    },
                    {
                        "role": "user",
                        "content": (
                            f"Query: {st.session_state.original_question}\nArticles:\n{articles_prompt}\n\n"
                            "For each article, provide a relevance score (0 to 1) as a JSON object mapping article IDs to scores."
                        ),
                    },
                ]
                with st.spinner("Filtering PubMed articles for question relevance"):
                    try:
                        response = create_chat_completion(messages, model="o3-mini")
                        relevance_scores = json.loads(response.choices[0].message.content.strip())
                        relevant_articles = [
                            article for article in filtered_articles
                            if float(relevance_scores.get(str(article["id"]), 0)) >= relevance_threshold
                        ]
                    except Exception as e:
                        logger.error(f"Error filtering articles: {e}")
                        relevant_articles = filtered_articles
                articles = [{
                    "id": a["id"],
                    "title": a["title"],
                    "link": a["link"],
                    "year": a["year"],
                    "abstract": a["abstract"],
                } for a in relevant_articles]
            else:
                articles = [{
                    "id": a["id"],
                    "title": a["title"],
                    "link": a["link"],
                    "year": a["year"],
                    "abstract": a["abstract"],
                } for a in filtered_articles]
            return articles[:max_results]
        except (aiohttp.ClientError, asyncio.TimeoutError) as e:
            print(f"Error fetching PubMed articles: {e}")
            return []

def realtime_search(query, domains, max, start_year=2020):
    url = "https://real-time-web-search.p.rapidapi.com/search"
    full_query = f"{query} AND ({domains})"
    start_date = f"{start_year}-01-01"
    end_date = datetime.now().strftime("%Y-%m-%d")
    querystring = {"q": full_query, "limit": max, "from": start_date, "to": end_date}
    headers = {
        "X-RapidAPI-Key": st.secrets["X-RapidAPI-Key"],
        "X-RapidAPI-Host": "real-time-web-search.p.rapidapi.com",
    }
    try:
        response = requests.get(url, headers=headers, params=querystring)
        response.raise_for_status()
        response_data = response.json().get("data", [])
        urls = [item.get("url") for item in response_data]
        snippets = [
            f"**{item.get('title')}**  \n*{item.get('snippet')}*  \n{item.get('url')} <END OF SITE>"
            for item in response_data
        ]
    except requests.exceptions.RequestException as e:
        st.error(f"RapidAPI real-time search failed to respond: {e}")
        return [], []
    return snippets, urls

#########################################
# OpenAI API Helper Functions for Chat Completions
#########################################
async def get_response(messages, model="o3-mini"):
    async with aiohttp.ClientSession() as session:
        payload = {"model": model, "messages": messages}
        if model != "o3-mini":
            payload["temperature"] = 0.3
        response = await session.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json=payload,
        )
        return await response.json()

async def get_responses(queries):
    tasks = [get_response(query) for query in queries]
    return await asyncio.gather(*tasks)

def clean_text(text):
    text = re.sub(r"([a-z])([A-Z])", r"\1 \2", text)
    text = text.replace("-", " ").replace(" .", ".")
    text = re.sub(r"\s{2,}", " ", text)
    return text

def refine_output(data):
    all_sources = ""
    for i, citation in enumerate(sorted(data, key=lambda x: x.get("score", 0), reverse=True)[:8], 1):
        normalized_score = round(citation.get("score", 0) * 100, 2)
        all_sources += f"**Source {i} (Relevance: {normalized_score}%)**\n\n"
        if "url" in citation:
            all_sources += f"[Link to source]({citation['url']})\n\n"
        cleaned_text = clean_text(citation.get("context", ""))
        truncated_text = cleaned_text[:3000] + "..." if len(cleaned_text) > 3000 else cleaned_text
        all_sources += f"{truncated_text}\n\n"
        if "Table" in cleaned_text:
            all_sources += "This source contained tabular data.\n\n"
        all_sources += "---\n\n"
    return all_sources

def get_db_path():
    return tempfile.mkdtemp(prefix="db_")

def extract_expert_info(json_input):
    data = json.loads(json_input)
    experts = []
    domains = []
    expert_questions = []
    for item in data["rephrased_questions"]:
        experts.append(item["expert"])
        domains.append(item["domain"])
        expert_questions.append(item["question"])
    return experts, domains, expert_questions

@st.cache_data
def create_chat_completion(
    messages,
    google=False,
    model="gpt-4o",
    frequency_penalty=0,
    logit_bias=None,
    logprobs=False,
    top_logprobs=None,
    max_completion_tokens=5000,
    n=1,
    presence_penalty=0,
    response_format=None,
    seed=None,
    stop=None,
    stream=False,
    include_usage=False,
    temperature=1,
    tools=None,
    tool_choice="none",
    user=None,
):
    if google:
        client = OpenAI(
            api_key=st.secrets["GOOGLE_API_KEY"],
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
        )
        params = {
            "model": model,
            "messages": messages,
            "stream": stream,
            "temperature": temperature,
        }
    else:
        client = OpenAI()
        params = {
            "model": model,
            "messages": messages,
            "frequency_penalty": frequency_penalty,
            "logit_bias": logit_bias,
            "logprobs": logprobs,
            "top_logprobs": top_logprobs,
            "max_completion_tokens": max_completion_tokens,
            "n": n,
            "presence_penalty": presence_penalty,
            "response_format": response_format,
            "seed": seed,
            "stop": stop,
            "stream": stream,
            "temperature": temperature,
            "user": user,
        }
    if model == "o3-mini":
        params.pop("temperature", None)
        params["reasoning_effort"] = "medium"
    if stream:
        params["stream_options"] = {"include_usage": include_usage}
    else:
        params.pop("stream_options", None)
    if tools:
        params["tools"] = [{"type": "function", "function": tool} for tool in tools]
        params["tool_choice"] = tool_choice
    if response_format == "json_object":
        params["response_format"] = {"type": "json_object"}
    elif response_format == "text":
        params["response_format"] = {"type": "text"}
    params = {k: v for k, v in params.items() if v is not None}
    completion = client.chat.completions.create(**params)
    return completion

def check_password() -> bool:
    if st.secrets["docker"] == "docker":
        st.session_state.password_correct = True
        return True
    if "password" not in st.session_state:
        st.session_state.password = ""
    if "password_correct" not in st.session_state:
        st.session_state.password_correct = False
    if "login_attempts" not in st.session_state:
        st.session_state.login_attempts = 0

    def password_entered() -> None:
        if st.session_state["password"] == st.secrets["password"]:
            st.session_state["password_correct"] = True
            st.session_state.login_attempts = 0
        else:
            st.session_state["password_correct"] = False
            st.session_state.login_attempts += 1

    if not st.session_state["password_correct"]:
        st.text_input("Password", type="password", on_change=password_entered, key="password")
        if st.session_state.login_attempts > 0:
            st.error(f"😕 Password incorrect. Attempts: {st.session_state.login_attempts}")
        st.write("*Please contact David Liebovitz, MD if you need an updated password for access.*")
        return False
    return True

#########################################
# Main Function: Orchestrates the App UI and Workflow
#########################################
def main():
    st.title("Helpful Answers with AI!")
    current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    db_path = get_db_path()

    # Configure the vectorstore app using our custom MyApp
    rag_model = "o3-mini"
    rag_provider = "openai"
    embedder_model = st.sidebar.toggle("Embedder Model: Use text-embedding-3-large", help="Toggle to use text-embedding-3-large.")
    embedder_model = "text-embedding-3-large" if embedder_model else "text-embedding-3-small"
    rag_key = api_key  # for simplicity in this snippet

    config = {
        "llm": {
            "provider": rag_provider,
            "config": {"model": rag_model, "stream": False, "api_key": rag_key},
        },
        "vectordb": {
            "provider": "chroma",
            "config": {"collection_name": "ai-helper", "dir": db_path, "allow_reset": True},
        },
        "embedder": {
            "provider": "openai",
            "config": {"api_key": api_key, "model": embedder_model},
        },
        "chunker": {
            "chunk_size": 5000,
            "chunk_overlap": 100,
            "length_function": "len",
            "min_chunk_size": 2000,
        },
    }

    # Use our custom MyApp in place of EmbedChain
    app = MyApp.from_config(config)
    
    with st.expander("About this app"):
        st.info(
            """This app interprets a user query and retrieves content from selected internet domains (including PubMed if applicable)
            for an initial answer. It then asks AI personas their opinions on the topic using updated content. The vectorstore is powered
            by Chroma and uses custom embedding and chunking logic. App author: David Liebovitz, MD."""
        )
    st.info("Please validate all guidance using the sources!")
    col1, col2 = st.columns([1, 1])
    if check_password():
        with col1:
            original_query = st.text_area("Ask a nice question...", placeholder="Enter your question here...", help="Ask any knowledge-based question.")
        st.session_state.original_question = original_query
        find_experts_messages = [
            {"role": "system", "content": system_prompt_expert_questions},
            {"role": "user", "content": original_query},
        ]
        determine_domain_messages = [
            {"role": "system", "content": choose_domain},
            {"role": "user", "content": original_query},
        ]
        first_view = False
        col2.write(" ")
        col2.write(" ")
        col2.write(" ")

        # Check if cutting-edge PubMed research should be included
        if st.sidebar.checkbox("Include Cutting-Edge Research in PubMed (default is consensus review articles)", help="Check to include latest, not yet consensus, articles.", value=False):
            pubmed_prompt = cutting_edge_pubmed_prompt
        else:
            pubmed_prompt = optimize_pubmed_search_terms_system_prompt

        deeper_dive = st.sidebar.checkbox("Deeper Dive", help="Check to include PubMed explicitly with extensive searching.", value=True)
        if col2.button("Begin Research"):
            first_view = True
            st.session_state.articles = []
            st.session_state.pubmed_search_terms = ""
            st.session_state.chosen_domain = ""
            st.session_state.followup_messages = []
            st.session_state.expert_answers = []
            st.session_state.messages1 = []
            st.session_state.messages2 = []
            st.session_state.messages3 = []
            st.session_state.initial_response_thread = []
            st.session_state.final_thread = []
            st.session_state.thread_with_tavily_context = []
            st.session_state.tavily_initial_response = []
            with col1:
                if deeper_dive:
                    with st.spinner("Determining the best domain for your question..."):
                        restrict_domains_response = create_chat_completion(determine_domain_messages, model="gpt-4o", temperature=0.3)
                        st.session_state.chosen_domain = restrict_domains_response.choices[0].message.content
                    if st.session_state.chosen_domain == "medical":
                        domains = medical_domains
                    else:
                        domains = st.session_state.chosen_domain
                    try:
                        app.reset()
                    except Exception:
                        st.error("Error resetting app; just proceed")
                    search_messages = [
                        {"role": "system", "content": optimize_search_terms_system_prompt},
                        {"role": "user", "content": f"considering it is {current_datetime}, {original_query}"},
                    ]
                    with st.spinner("Optimizing search terms..."):
                        try:
                            response_google_search_terms = create_chat_completion(search_messages, temperature=0.3)
                        except Exception as e:
                            st.error(f"Error during OpenAI call: {e}")
                    google_search_terms = response_google_search_terms.choices[0].message.content
                    st.session_state.chosen_domain = st.session_state.chosen_domain.replace('"', "").replace("'", "")
                    if st.session_state.chosen_domain == "medical":
                        pubmed_messages = [{"role": "system", "content": pubmed_prompt}, {"role": "user", "content": original_query}]
                        response_pubmed_search_terms = create_chat_completion(pubmed_messages, temperature=0.3)
                        pubmed_search_terms = response_pubmed_search_terms.choices[0].message.content
                        st.session_state.pubmed_search_terms = pubmed_search_terms
                        with st.spinner(f'Searching PubMed for "{pubmed_search_terms}"...'):
                            articles = asyncio.run(pubmed_abstracts(pubmed_search_terms))
                        st.session_state.articles = articles
                        with st.spinner("Adding PubMed abstracts to the knowledge base..."):
                            if articles:
                                for article in articles:
                                    retries = 3
                                    success = False
                                    while retries > 0 and not success:
                                        try:
                                            link = article.get("link")
                                            if not link:
                                                raise ValueError("Article missing link")
                                            app.add(link, data_type="web_page")
                                            success = True
                                        except Exception as e:
                                            retries -= 1
                                            logger.error(f"Error adding article {article}: {str(e)}")
                                            if retries > 0:
                                                time.sleep(1)
                                            else:
                                                st.error("PubMed results did not meet relevance and recency check.")
                        
                                st.write("Documents from PubMed in collection:", app.collection.peek())

                        
                        if not articles:
                            st.warning("No recent and relevant PubMed articles identified for the knowledge base.")
                        else:
                            with st.expander("View PubMed Results While Waiting"):
                                pubmed_link = "https://pubmed.ncbi.nlm.nih.gov/?term=" + st.session_state.pubmed_search_terms
                                st.page_link(pubmed_link, label="Click here to view in PubMed", icon="📚")
                                for article in articles:
                                    st.markdown(f"### [{article['title']}]({article['link']})")
                                    st.write(f"Year: {article['year']}")
                                    st.write(article["abstract"] if article["abstract"] else "No abstract available")
                    with st.spinner(f'Searching for "{google_search_terms}"...'):
                        if st.sidebar.radio("Internet search provider:", options=["Google", "Exa"]) == "Google":
                            st.session_state.snippets, st.session_state.urls = realtime_search(google_search_terms, domains, st.sidebar.number_input("Number of web pages to retrieve:", min_value=1, max_value=20, value=6))
                        else:
                            # Using Exa as fallback (if configured)
                            three_years_ago = datetime.now() - timedelta(days=3 * 365.25)
                            date_cutoff = three_years_ago.strftime("%Y-%m-%d")
                            search_response = exa.search_and_contents(
                                google_search_terms,
                                text={"include_html_tags": False, "max_characters": 1000},
                                highlights={"highlights_per_url": 2, "num_sentences": 5, "query": "This is the highlight query:"},
                                start_published_date=date_cutoff,
                            )
                            st.session_state.snippets = [result.text for result in search_response.results]
                            st.session_state.urls = [result.url for result in search_response.results]
                    with st.expander("View Reliable Internet Results While Waiting"):
                        for snippet in st.session_state.snippets:
                            st.markdown(snippet.replace("<END OF SITE>", ""))
                    blocked_sites = []
                    with st.spinner("Retrieving full content from web pages..."):
                        for site in st.session_state.urls:
                            try:
                                app.add(site, data_type="web_page")
                            except Exception:
                                blocked_sites.append(site)
                        st.write("Internet added to collection:", app.collection.peek())

                    with st.spinner("Analyzing retrieved content..."):
                        try:
                            prepare_rag_query_messages = [
                                {"role": "system", "content": prepare_rag_query},
                                {"role": "user", "content": f"User query to refine: {original_query}"},
                            ]
                            query_for_rag = create_chat_completion(prepare_rag_query_messages, model="gpt-4o-mini", temperature=0.3)
                            updated_rag_query = query_for_rag.choices[0].message.content
                        except Exception as e:
                            st.error(f"Error during rag prep {e}")
                        try:
                            citations = app.search(updated_rag_query, num_documents=10  )
                            filtered = filter_citations(citations)
                            st.session_state.citations = filtered
                        except Exception as e:
                            st.error(f"Error during semantic query: {e}")
                        try:
                            updated_answer_prompt = rag_prompt2.format(
                                question=original_query,
                                context=st.session_state.citations,
                            )
                            prepare_updated_answer_messages = [{"role": "user", "content": updated_answer_prompt}]
                            updated_answer = create_chat_completion(prepare_updated_answer_messages, model="gpt-4o-mini", temperature=0.3)
                            updated_answer_text = updated_answer.choices[0].message.content
                        except Exception as e:
                            st.error(f"Error during second pass: {e}")
                        if updated_answer is not None:
                            st.session_state.full_initial_response = updated_answer_text
                    container1 = st.container()
                    with container1:
                        st.info(f"Response as of **{current_datetime}:**\n\n")
                        st.markdown(st.session_state.full_initial_response)
                        display_url_list(st.session_state.citations)
                        with st.expander("View Sources"):
                            display_citations(st.session_state.citations)
                        if st.button("Create Word Document"):
                            doc = markdown_to_word(st.session_state.rag_response)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word Document",
                                data=buffer,
                                file_name="prelim_response.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                else:
                    # Quick search path (using Tavily)
                    with st.spinner("Determining the best domain for your question..."):
                        restrict_domains_response = create_chat_completion(determine_domain_messages, model="gpt-4o", temperature=0.3)
                        st.session_state.chosen_domain = restrict_domains_response.choices[0].message.content
                    tavily_initial_search_domains = tavily_domains if st.session_state.chosen_domain == "medical" else ""
                    try:
                        tavily_client = TavilyClient(api_key=st.secrets["TAVILY_API_KEY"])
                    except Exception as e:
                        st.error(f"Error during Tavily client initialization: {e}")
                        return
                    with st.spinner("Accessing reliable internet domains for updates..."):
                        try:
                            updated_initial_tavily_query = create_chat_completion([
                                {"role": "system", "content": "Optimize the user question for submission to an online search engine. Return only the optimized question for use in a python pipeline."},
                                {"role": "user", "content": original_query},
                            ])
                            response = tavily_client.search(
                                query=updated_initial_tavily_query.choices[0].message.content,
                                include_domains=tavily_initial_search_domains,
                                search_depth="advanced",
                            )
                        except json.JSONDecodeError as e:
                            print(f"Error decoding JSON: {e}")
                        st.session_state.tavily_urls = extract_and_format_urls(response)
                        results = response.get("results", [])
                        result_texts = [
                            f"**Title:** {result['title']}\n\n**URL:** {result['url']}\n\n**Content:** {result['content']}\n\n**Relevancy Score:** {result['score']:.2f}\n\n"
                            for result in results
                        ]
                        st.session_state.tavily_initial_response = "\n".join(result_texts)
                    updated_initial_question_with_tavily = original_query + f" If appropriate, incorporate these updated findings or references from reliable sites: {st.session_state.tavily_initial_response}"
                    st.session_state.initial_response_thread.append({"role": "user", "content": updated_initial_question_with_tavily})
                    try:
                        updated_answer = create_chat_completion(st.session_state.initial_response_thread, model="o3-mini")
                    except Exception as e:
                        st.error(f"Error during second pass: {e}")
                    if updated_answer:
                        current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        full_response = f"As of **{current_datetime}**:\n\n" + updated_answer.choices[0].message.content
                        st.session_state.full_initial_response = full_response
                        first_view = True
                container1 = col1.container()
                with container1:
                    if st.session_state.tavily_initial_response:
                        with st.expander("View Tavily Initial Search Results"):
                            st.write(st.session_state.tavily_initial_response)
                        st.info("Initial Response")
                        st.markdown(st.session_state.full_initial_response)
                        st.markdown(st.session_state.tavily_urls)
                        if st.button("Create Word Document for Quick Search"):
                            doc = markdown_to_word(st.session_state.full_initial_response)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word Document for Quick Search",
                                data=buffer,
                                file_name="full_initial_response.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
        with col2:
            if st.session_state.rag_response or st.session_state.full_initial_response:
                initial_followup = st.checkbox("Ask Follow-Up Questions for Initial Response")
                if initial_followup:
                    add_internet_content = True
                    formatted_output = []
                    for citation in st.session_state.citations:
                        try:
                            source_metadata, source_text = citation, citation.get("context", "")
                            metadata_details = "\n".join(f"{key.capitalize()}: {value}" for key, value in source_metadata.items())
                            formatted_output.append(f"{metadata_details}\nSource text: {source_text}\n---")
                        except Exception as e:
                            formatted_output.append(f"Error processing citation: {citation}\nError: {str(e)}\n---")
                    formatted_output_str = "\n".join(formatted_output)
                    prelim_followup_prompt2 = prelim_followup_prompt.format(
                        prior_question=original_query,
                        evidence=formatted_output_str,
                        prior_answer=st.session_state.full_initial_response,
                    )
                    if not st.session_state.initial_response_thread:
                        st.session_state.initial_response_thread.append({"role": "system", "content": prelim_followup_prompt2})
                    if initial_followup_question := st.chat_input("Ask followup!"):
                        if add_internet_content:
                            try:
                                tavily_client = TavilyClient(api_key=st.secrets["TAVILY_API_KEY"])
                            except Exception as e:
                                st.error(f"Error during Tavily client initialization: {e}")
                                return
                            with st.spinner("Retrieving additional internet content..."):
                                try:
                                    updated_tavily_query = create_chat_completion([
                                        {"role": "system", "content": "Combine user inputs (an initial and then followup question) into one question optimized for searching online. Return only the optimized question."},
                                        {"role": "user", "content": f"{original_query} and {initial_followup_question}"},
                                    ])
                                    response = tavily_client.search(
                                        query=updated_tavily_query.choices[0].message.content,
                                        include_domains=tavily_domains,
                                        search_depth="advanced",
                                    )
                                except json.JSONDecodeError as e:
                                    print(f"Error decoding JSON: {e}")
                                results = response.get("results", [])
                                result_texts = [
                                    f"**Title:** {result['title']}\n\n**URL:** {result['url']}\n\n**Content:** {result['content']}\n\n**Relevancy Score:** {result['score']:.2f}\n\n"
                                    for result in results
                                ]
                                st.session_state.tavily_followup_response = "\n".join(result_texts)
                            with st.expander("New Retrieved Search Content"):
                                st.write(f"Updated Query: {updated_tavily_query.choices[0].message.content}")
                                st.write("\n\n")
                                st.write(st.session_state.tavily_followup_response)
                            updated_followup_question = initial_followup_question + f" Here's more of what I found online: {st.session_state.tavily_followup_response}"
                            st.session_state.thread_with_tavily_context = st.session_state.initial_response_thread
                            st.session_state.thread_with_tavily_context.append({"role": "user", "content": updated_followup_question})
                            initial_followup_messages = st.session_state.thread_with_tavily_context
                        else:
                            st.session_state.initial_response_thread.append({"role": "user", "content": initial_followup_question})
                            initial_followup_messages = st.session_state.initial_response_thread
                        with st.chat_message("user"):
                            st.markdown(initial_followup_question)
                        with st.chat_message("assistant"):
                            client = OpenAI()
                            try:
                                stream = client.chat.completions.create(
                                    model="o3-mini",
                                    messages=[{"role": m["role"], "content": m["content"]} for m in initial_followup_messages],
                                    stream=True,
                                )
                                response = st.write_stream(stream)
                            except Exception as e:
                                st.error(f"Error during OpenAI call: {e}")
                                return
                            st.session_state.initial_response_thread.append({"role": "assistant", "content": response})
                    with st.expander("View full follow-up thread"):
                        for message in st.session_state.initial_response_thread:
                            if message["role"] != "system":
                                emoji = role_emojis.get(message["role"], "❓")
                                st.write(f"{emoji} {message['role'].capitalize()}: {message['content']}")
                    if st.session_state.initial_response_thread:
                        if st.checkbox("Download Followup Conversation"):
                            full_followup_conversation = ""
                            for message in st.session_state.initial_response_thread:
                                if message["role"] != "system":
                                    emoji = role_emojis.get(message["role"], "❓")
                                    full_followup_conversation += f"{emoji} {message['role'].capitalize()}: {message['content']}\n\n"
                            html = markdown2.markdown(full_followup_conversation, extras=["tables"])
                            st.download_button("Download Followup Conversation", html, "followup_conversation.html", "text/html")
                if not initial_followup:
                    if st.button("Ask 3 AI Expert Personas for Opinions"):
                        prelim_response = st.session_state.full_initial_response + str(st.session_state.citations)
                        try:
                            completion = create_chat_completion(messages=find_experts_messages, model="o3-mini", temperature=0.3, response_format="json_object")
                        except Exception as e:
                            st.error(f"Error during OpenAI call: {e}")
                            return
                        json_output = completion.choices[0].message.content
                        experts, domains, expert_questions = extract_expert_info(json_output)
                        st.session_state.experts = experts
                        updated_expert1_system_prompt = expert1_system_prompt.format(expert=experts[0], domain=domains[0])
                        updated_expert2_system_prompt = expert2_system_prompt.format(expert=experts[1], domain=domains[1])
                        updated_expert3_system_prompt = expert3_system_prompt.format(expert=experts[2], domain=domains[2])
                        updated_question1 = expert_questions[0]
                        updated_question2 = expert_questions[1]
                        updated_question3 = expert_questions[2]
                        prelim_response = st.session_state.rag_response + st.session_state.source_chunks
                        expert1_messages = [
                            {"role": "system", "content": updated_expert1_system_prompt},
                            {"role": "user", "content": updated_question1 + "Here's what I already found online: " + prelim_response},
                        ]
                        st.session_state.messages1 = expert1_messages
                        expert2_messages = [
                            {"role": "system", "content": updated_expert2_system_prompt},
                            {"role": "user", "content": updated_question2 + "Here's what I already found online: " + prelim_response},
                        ]
                        st.session_state.messages2 = expert2_messages
                        expert3_messages = [
                            {"role": "system", "content": updated_expert3_system_prompt},
                            {"role": "user", "content": updated_question3 + "Here's what I already found online: " + prelim_response},
                        ]
                        st.session_state.messages3 = expert3_messages
                        with st.spinner("Waiting for experts to respond..."):
                            st.session_state.expert_answers = asyncio.run(get_responses([expert1_messages, expert2_messages, expert3_messages]))
        with col1:
            if not first_view:
                try:
                    if st.session_state.pubmed_search_terms and st.session_state.articles:
                        with st.expander("View PubMed Results Added to Knowledge Base"):
                            pubmed_link = "https://pubmed.ncbi.nlm.nih.gov/?term=" + st.session_state.pubmed_search_terms
                            st.page_link(pubmed_link, label="Click here to view in PubMed", icon="📚")
                            for article in st.session_state.articles:
                                st.markdown(f"### [{article['title']}]({article['link']})")
                                st.write(f"Year: {article['year']}")
                                st.write(article["abstract"] if article["abstract"] else "No abstract available")
                except:
                    st.write("No Relevant PubMed articles to display - if topic works, API may be down!")
                    pubmed_link = "https://pubmed.ncbi.nlm.nih.gov/?term=" + st.session_state.pubmed_search_terms
                    st.page_link(pubmed_link, label="Click here to try in PubMed", icon="📚")
                with st.expander("View Internet Results Added to Knowledge Base"):
                    for snippet in st.session_state.snippets:
                        st.markdown(snippet.replace("<END OF SITE>", ""))
                if st.session_state.full_initial_response:
                    container1 = st.container()
                    with container1:
                        st.info(f"Response as of **{current_datetime}:**\n\n")
                        st.markdown(st.session_state.full_initial_response)
                        if st.session_state.citations:
                            display_url_list(st.session_state.citations)
                            with st.expander("View Source Excerpts"):
                                display_citations(st.session_state.citations)
                        if st.session_state.tavily_initial_response:
                            st.markdown(st.session_state.tavily_urls)
                            with st.expander("Retrieved Search Content"):
                                st.write(st.session_state.tavily_initial_response)
                        if st.button("Create Word File"):
                            doc = markdown_to_word(st.session_state.full_initial_response)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File",
                                data=buffer,
                                file_name="prelim_response.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
        with col2:
            if st.session_state.expert_answers:
                container2 = st.container()
                with container2:
                    st.info("AI Expert Persona Responses")
                    with st.expander(f"AI {st.session_state.experts[0]} Perspective"):
                        expert_0 = st.session_state.expert_answers[0]["choices"][0]["message"]["content"]
                        st.write(expert_0)
                        st.session_state.messages1.append({"role": "assistant", "content": expert_0})
                        if st.button("Create Word File for AI Expert 1"):
                            doc = markdown_to_word(expert_0)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File for AI Expert 1",
                                data=buffer,
                                file_name="AI_expert_1.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                    with st.expander(f"AI {st.session_state.experts[1]} Perspective"):
                        expert_1 = st.session_state.expert_answers[1]["choices"][0]["message"]["content"]
                        st.write(expert_1)
                        st.session_state.messages2.append({"role": "assistant", "content": expert_1})
                        if st.button("Create Word File for AI Expert 2"):
                            doc = markdown_to_word(expert_1)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File for AI Expert 2",
                                data=buffer,
                                file_name="AI_expert_2.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                    with st.expander(f"AI {st.session_state.experts[2]} Perspective"):
                        expert_2 = st.session_state.expert_answers[2]["choices"][0]["message"]["content"]
                        st.write(expert_2)
                        st.session_state.messages3.append({"role": "assistant", "content": expert_2})
                        if st.button("Create Word File for AI Expert 3"):
                            doc = markdown_to_word(expert_2)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File for AI Expert 3",
                                data=buffer,
                                file_name="AI_expert_3.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )

if __name__ == "__main__":
    main()
