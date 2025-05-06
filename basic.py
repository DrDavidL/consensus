#########################################
# Import Libraries and Setup
#########################################
import asyncio
import json
import re
import os
import tempfile
import time
import warnings

# Skip alembic patching - not needed with current version
def skip_alembic_init():
    """
    Function to skip alembic initialization in EmbedChain.
    This will be implemented in the App initialization instead.
    """
    pass

# Suppress specific deprecation and syntax warnings
warnings.filterwarnings("ignore", category=SyntaxWarning)
warnings.filterwarnings("ignore", category=UserWarning, module="langchain")
warnings.filterwarnings("ignore", category=DeprecationWarning, message="Testing an element's truth value")
warnings.filterwarnings("ignore", message="Accessing the 'model_fields' attribute on the instance is deprecated")
warnings.filterwarnings("ignore", message="Accessing the 'model_fields' attribute on the instance is deprecated. Instead, you should access this attribute from the model class. Deprecated in Pydantic V2.11 to be removed in V3.0.")
# Remove the following line, as KeyError is not a Warning subclass and will cause an error
# warnings.filterwarnings("ignore", category=KeyError, message="'script'")
from datetime import datetime, timedelta
from typing import List, Tuple, Dict
import xml.etree.ElementTree as ET
from io import BytesIO
from docx.shared import Pt

#########################################
# Third-Party Library Imports
#########################################
import aiohttp
import requests
import streamlit as st
import anthropic
from openai import OpenAI
from exa_py import Exa
import markdown2
from docx import Document
from tavily import TavilyClient

from ragas import SingleTurnSample
from ragas.metrics import Faithfulness
from ragas.metrics import AspectCritic
from ragas.metrics import RubricsScore
from ragas.llms import LangchainLLMWrapper
from ragas.embeddings import LangchainEmbeddingsWrapper
from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings
evaluator_llm = LangchainLLMWrapper(ChatOpenAI(model="gpt-4.1"))
evaluator_embeddings = LangchainEmbeddingsWrapper(OpenAIEmbeddings())

#########################################
# EmbedChain Library Imports
#########################################
from embedchain import App
from embedchain.config import BaseLlmConfig

#########################################
# Logging Configuration
#########################################
import logging

#########################################
# Global Variables and Logging Configuration
#########################################

# Standard timeouts for network requests
AIOHTTP_TIMEOUT = aiohttp.ClientTimeout(total=60)  # 60 seconds total for aiohttp requests
REQUESTS_TIMEOUT = 60  # 60 seconds for requests library

role_emojis = {
    "user": "ðŸ‘¤",
    "assistant": "ðŸ¤–",
}

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

NAME_PATTERN = re.compile(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?$')
AUTHOR_ET_AL_PATTERN = re.compile(r'^[A-Z][a-z]+\s+et\s+al\.?$', re.IGNORECASE)
NUMBERED_REF_PATTERN = re.compile(r'^\d+\.\s', re.MULTILINE)

PUBMED_REGEX = re.compile(r'PubMed:', re.IGNORECASE)
PMC_REGEX = re.compile(r'PMC', re.IGNORECASE)
DOI_PATTERN = re.compile(r'\bdoi:\s*\S+', re.IGNORECASE)
PMID_PATTERN = re.compile(r'\bPMID:\s*\d+')
MONTH_PATTERN = re.compile(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b')
LOWER_UPPER_PATTERN = re.compile(r'([a-z])([A-Z])')
NUMBERED_REF_PATTERN = re.compile(r'^\d+\.\s', re.MULTILINE)

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if color is not None:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

#########################################
# Import Prompts for AI Guidance and Search
#########################################
from prompts import (
    system_prompt_expert_questions,
    expert1_system_prompt,
    expert2_system_prompt,
    expert3_system_prompt,
    optimize_search_terms_system_prompt,
    optimize_pubmed_search_terms_system_prompt,
    cutting_edge_pubmed_prompt,
    prepare_rag_query,
    rag_prompt2,
    choose_domain,
    medical_domains,
    prelim_followup_prompt,
    tavily_domains,
)

#########################################
# Streamlit App Configuration and API Keys Setup
#########################################
st.set_page_config(
    page_title="Research Help with AI! ðŸš€",
    layout="wide",
    page_icon="âœ¨",
    initial_sidebar_state="expanded",
)
api_key = st.secrets["OPENAI_API_KEY"]
api_key_anthropic = st.secrets["ANTHROPIC_API_KEY"]
exa = Exa(st.secrets["EXA_API_KEY"])

#########################################
# Session State Initialization
#########################################
if "snippets" not in st.session_state:
    st.session_state["snippets"] = []
if "urls" not in st.session_state:
    st.session_state["urls"] = []
if "rag_response" not in st.session_state:
    st.session_state["rag_response"] = ""
if "source_chunks" not in st.session_state:
    st.session_state.source_chunks = ""
if "experts" not in st.session_state:
    st.session_state.experts = []
if "messages1" not in st.session_state:
    st.session_state.messages1 = []
if "messages2" not in st.session_state:
    st.session_state.messages2 = []
if "messages3" not in st.session_state:
    st.session_state.messages3 = []
if "followup_messages" not in st.session_state:
    st.session_state.followup_messages = []
if "final_thread" not in st.session_state:
    st.session_state.final_thread = []
if "expert_number" not in st.session_state:
    st.session_state.expert_number = 0
if "expert_answers" not in st.session_state:
    st.session_state.expert_answers = []
if "original_question" not in st.session_state:
    st.session_state.original_question = ""
if "chosen_domain" not in st.session_state:
    st.session_state.chosen_domain = ""
if "articles" not in st.session_state:
    st.session_state.articles = []
if "pubmed_search_terms" not in st.session_state:
    st.session_state.pubmed_search_terms = ""
if "full_initial_response" not in st.session_state:
    st.session_state.full_initial_response = ""
if "initial_response_thread" not in st.session_state:
    st.session_state.initial_response_thread = []
if "citations" not in st.session_state:
    st.session_state.citations = []
if "thread_with_tavily_context" not in st.session_state:
    st.session_state.thread_with_tavily_context = []
if "tavily_followup_response" not in st.session_state:
    st.session_state.tavily_followup_response = ""
if "tavily_initial_response" not in st.session_state:
    st.session_state.tavily_initial_response = ""
if "tavily_urls" not in st.session_state:
    st.session_state.tavily_urls = ""
if "ragas_score" not in st.session_state:
    st.session_state.ragas_score = 0.0
if "older_pubmed_articles_alert" not in st.session_state:
    st.session_state.older_pubmed_articles_alert = False

#########################################
# Sidebar Configuration: UI Elements & Settings
#########################################
with st.sidebar:
    st.title("Main Settings")
    if "short_use_case" not in st.session_state:
        st.session_state.short_use_case = "Answer the Question"
    short_use_case = st.radio("Use Case", ["Answer the Question", "Helpful PubMed Query", "Helpful Internet Sites"], key="short_use_case")

    if short_use_case == "Helpful PubMed Query":
        st.info("Use this option to generate an advanced PubMed query.")
    elif short_use_case == "Helpful Internet Sites":
        st.info("Use this option for a list of reliable webites to answer your question.")
    elif short_use_case == "Model Guidance Only":
        st.info("Use this option to get model guidance on a question.")
        st.write("Select the model guidance settings below.")




    st.divider()
    with st.sidebar.expander("Advanced Settings"):
        st.info(
            "Default settings are fine for most use cases!"
        )
        # Toggle for subject area model
        topic_model_choice = st.toggle(
            "Subject Area: Use GPT-4.1",
            help="Toggle to use GPT-4.1 model for determining if medical; otherwise, 4o-mini.",
        )
        if topic_model_choice:
            st.write("GPT-4.1 model selected.")
            topic_model = "gpt-4.1"
        else:
            st.write("GPT-4o-mini model selected.")
            topic_model = "gpt-4o-mini"
        st.divider()

        # Web search settings
        search_type = "all"

        site_number = st.number_input(
            "Number of web pages to retrieve:",
            min_value=1,
            max_value=20,
            value=6,
            step=1,
        )
        internet_search_provider = st.radio(
            "Internet search provider:",
            options=["Google", "Exa"],
            horizontal=True,
            help="Only specific Google domains are used for retrieving current Medical or General Knowledge. Exa.ai is a new type of search tool that predicts relevant sites; domain filtering not yet added here.",
        )
        if internet_search_provider == "Google":
            st.info("Web domains used for medical questions.")
            edited_medical_domains = st.text_area(
                "Edit domains (maintain format pattern):", medical_domains, height=200
            )
        st.divider()

        # PubMed search settings
        years_back = st.slider(
            "Years Back for PubMed Search",
            min_value=1,
            max_value=10,
            value=4,
            step=1,
            help="Set the number of years back to search PubMed.",
        )
        st.divider()
        max_results = st.slider(
            "Number of Abstracts to Review",
            min_value=3,
            max_value=20,
            value=6,
            step=1,
            help="Set the number of abstracts to review.",
        )
        st.divider()
        filter_relevance = st.toggle(
            "Filter Relevance of PubMed searching", value=True, help="Toggle to deselect."
        )
        if filter_relevance:
            relevance_threshold = st.slider(
                "Relevance Threshold",
                min_value=0.3,
                max_value=1.0,
                value=0.7,
                step=0.05,
                help="Set the minimum relevance score to consider an item relevant.",
            )
        else:
            relevance_threshold = 0.65
            st.write("Top sources will be added to the database regardless.")
        st.divider()

        # Technical settings for embedder model
        st.divider()
        embedder_model_choice = st.radio(
            "Embedder Model Options",
            ["text-embedding-3-small", "text-embedding-3-large", "gemini-embedding-exp-03-07"],
            index=1,
            help="Select the embedder model to use for the AI responses.",
        )
        if embedder_model_choice == "text-embedding-3-small":
            st.write("text-embedding-3-small model selected.")
            embedder_model = "text-embedding-3-small"
            embedder_provider = "openai"
            embedder_api_key = api_key
        elif embedder_model_choice == "text-embedding-3-large":
            st.write("text-embedding-3-large model selected.")
            embedder_model = "text-embedding-3-large"
            embedder_provider = "openai"
            embedder_api_key = api_key
        elif embedder_model_choice == "gemini-embedding-exp-03-07":
            st.write("gemini-embedding-exp-03-07 model selected.")
            embedder_model = "models/gemini-embedding-exp-03-07"
            embedder_provider = "google"
            os.environ["GOOGLE_API_KEY"] = st.secrets["GOOGLE_API_KEY"]
        # st.divider()
        # st.info(
        #     "GPT-4o-mini performs well for other options. For more complex synthesis, stay with GPT-4o or use Claude-3.5 Sonnet."
        # )

        # RAG model options
        rag_model_choice = "GPT-4.1"
        # rag_model_choice = st.radio(
        #     "RAG Model Options",
        #     ["GPT-4o-mini", "GPT-4o", "Claude-3.5 Sonnet", "Gemini-2"],
        #     index=1,
        #     help="Select the RAG model to use for the AI responses.",
        # )   
        if rag_model_choice == "GPT-4.1":
            # st.write("GPT-4o model selected.")
            rag_model = "gpt-4.1"
            rag_provider = "openai"
            rag_key = api_key
        elif rag_model_choice == "Gemini-2":
            st.write("Gemini-2 flash model selected.")
            rag_model = "gemini-2.0-flash"
            rag_provider = "google"
            rag_key = st.secrets["GOOGLE_API_KEY"]
        elif rag_model_choice == "GPT-4o-mini":
            st.write("GPT-4o-mini model selected.")
            rag_model = "gpt-4o-mini"
            rag_provider = "openai"
            rag_key = api_key
        elif rag_model_choice == "Claude-3.5 Sonnet":
            st.write("Claude-3-5-sonnet-latest model selected.")
            rag_model = "claude-3-5-sonnet-latest"
            rag_provider = "anthropic"
            rag_key = api_key_anthropic
        st.divider()

        # Second review model options
        second_review_model = st.radio(
            "Content Augmented Model Options",
            ["GPT-4o-mini", "GPT-4o", "GPT-4.1", "o3-mini", "Claude-3.7 Sonnet", "Gemini-2", "Gemini-2.5-flash", "Gemini-2.5-pro"],
            index=6,
            help="Select the RAG model to use for the AI responses.",
        )
        model_map = {
            "GPT-4o": {
                "message": "GPT-4o model selected.",
                "model": "gpt-4o",
                "provider": "openai",
                "key": api_key
            },
            "GPT-4.1": {
                "message": "GPT-4.1 model selected.",
                "model": "gpt-4.1",  
                "provider": "openai",
                "key": api_key
            },
            "Claude-3.7 Sonnet": {
                "message": "Claude-3-7-sonnet model selected.",
                "model": "claude-3-7-sonnet-20250219",
                "provider": "anthropic",
                "key": api_key_anthropic
            },
            "GPT-4o-mini": {
                "message": "GPT-4o-mini model selected.",
                "model": "gpt-4o-mini",
                "provider": "openai",
                "key": api_key
            },
            "Gemini-2": {
                "message": "Gemini-2 flash model selected.",
                "model": "gemini-2.0-flash",
                "provider": "google",
                "key": st.secrets["GOOGLE_API_KEY"]
            },
            "Gemini-2.5-flash": {
                "message": "Gemini-2.5 flash model selected.",
                "model": "gemini-2.5-flash-preview-04-17",
                "provider": "google",
                "key": st.secrets["GOOGLE_API_KEY"]
            },
            "Gemini-2.5-pro": {
                "message": "Gemini-2.5 pro model selected.",
                "model": "gemini-2.5-pro-preview-03-25",
                "provider": "google",
                "key": st.secrets["GOOGLE_API_KEY"]
            },
            "o3-mini": {
                "message": "o3-mini reasoning model selected.",
                "model": "o3-mini",
                "provider": "openai",
                "key": api_key
            }
        }

        if second_review_model in model_map:
            config = model_map[second_review_model]
            st.write(config["message"])
            second_model = config["model"]
            second_provider = config["provider"]
            second_key = config["key"]

        st.divider()

        # Expert personas model choice
        experts_model_choice = st.toggle(
            "3 AI Experts Model: Use GPT-4.1",
            help="Toggle to use GPT-4.1 model for expert responses; otherwise, o3-mini with reasoning.",
        )
        if experts_model_choice:
            st.write("GPT-4.1 model selected.")
            experts_model = "gpt-4.1"
        else:
            st.write("o3-mini reasoning model selected.")
            experts_model = "o3-mini"

        # Check if cutting-edge PubMed research should be included
        cutting_edge = st.checkbox(
            "Include Cutting-Edge Research in PubMed (default is consensus review articles)",
            help="Check to include latest, not yet consensus, articles in the search for medical content.",
            value=False,
        )
        if cutting_edge:
            pubmed_prompt = cutting_edge_pubmed_prompt
        else:
            pubmed_prompt = optimize_pubmed_search_terms_system_prompt

        deeper_dive = st.checkbox(
            "Deeper Dive",
            help="Check to include PubMed explicitly with extensive searching.",
            value=True,
        )

#########################################
# Utility Functions
#########################################


def is_non_informative(context: str) -> bool:
    """
    Returns True if the citation context is deemed non-informative.
    The following criteria are used:
      1. The context is empty or very short.
      2. The context appears to be a simple author name or follows a basic name pattern.
      3. The context appears to be a publication reference block 
         (e.g., includes markers like "PMID:" or "doi:" multiple times, numbered references,
          or repeated mentions of "PubMed:"/ "PMC").
      4. The context appears to be a disclosure or print prompt block.
    """
    context = context.strip()
    
    # Criterion 1: Empty or very short context
    if not context or len(context) < 5:
        return True

    # Criterion 2: Simple name patterns
    if NAME_PATTERN.fullmatch(context):
        return True

    if AUTHOR_ET_AL_PATTERN.fullmatch(context):
        return True

    # New Criterion: Check if combined PubMed and PMC mentions exceed 5.
    pubmed_mentions = len(PUBMED_REGEX.findall(context))
    pmc_mentions = len(PMC_REGEX.findall(context))
    if (pubmed_mentions + pmc_mentions) > 5:
        return True

    # Criterion 3: Check for publication reference block characteristics
    doi_matches = DOI_PATTERN.findall(context)
    pmid_matches = PMID_PATTERN.findall(context)
    
    # Density check for DOIs/PMIDs
    if (len(doi_matches) + len(pmid_matches)) > 3:
        return True

    # Check for reference numbering pattern (e.g., "47." or "1.")
    if len(NUMBERED_REF_PATTERN.findall(context)) > 3:
        return True

    # Check for month abbreviations in longer contexts with multiple "et al" mentions
    if context.count("et al") > 1 or MONTH_PATTERN.search(context):
        if len(context) > 150:
            return True

    # Criterion 4: Disclosure or print prompt blocks
    if context.count("Disclosure:") > 1:
        return True

    if "Print this section" in context or "What would you like to print?" in context:
        return True

    # Additional check: High line-break density typical in reference blocks
    if context.count('\n') > 5:
        return True

    # Additional check: Overly long contexts with common website cues
    if len(context) > 500 and ("Medscape" in context or "Copyright" in context):
        return True

    return False

def filter_citations(citations: list) -> list:
    """
    Given a list of citation dictionaries, filter out those entries where the 'context'
    is deemed non-informative. URLs are retained as part of the final dictionary.
    """
    return [
        {
            "context": citation.get("context", ""),
            "url": citation.get("metadata", {}).get("url", ""),
            "score": citation.get("metadata", {}).get("score", 0),
        }
        for citation in citations
        if not is_non_informative(citation.get("context", ""))
    ]

def extract_and_format_urls(tavily_output):
    """
    Extracts all URLs from the Tavily output and returns a formatted string
    with a numbered list. Each URL is printed as the link text.
    """
    results = tavily_output.get("results", [])
    if not results:
        return "No URLs found."

    # Extract URLs from each result (ignoring any missing or empty values)
    urls = [result.get("url", "") for result in results if result.get("url", "")]
    # Optionally remove duplicates (if necessary)
    unique_urls = sorted(set(urls))

    # Create a nicely formatted string with a numbered list of URLs.
    output_lines = ["List of References:"]
    for idx, url in enumerate(unique_urls, start=1):
        output_lines.append(f"{idx}. {url}")

    return "\n".join(output_lines)


def display_url_list(citations):
    """
    Extract and display a deduplicated list of URLs from the given citations.

    Each citation is expected to be a dictionary with at least a 'url' key.
    The function creates a clickable Markdown list where the URL itself is used as the link text.
    """
    # Extract URLs from each citation if present.
    urls = {citation.get("url", "") for citation in citations if citation.get("url", "")}

    # Remove duplicate URLs and sort for consistency.
    unique_urls = sorted(urls)

    # Display header.
    st.markdown("**List of Source URLs**", unsafe_allow_html=True)

    # Loop through the deduplicated URLs and display them as clickable markdown links.
    for url in unique_urls:
        # The markdown syntax [url](url) displays the URL as both link text and destination.
        st.markdown(f"- [{url}]({url})", unsafe_allow_html=True)


def display_citations(citations):
    """
    Display citations nicely in a Streamlit app.

    Each citation in the citations list is a dictionary with keys:
        - context (str): The citation text snippet.
        - url (str): The URL of the source.
        - score (float): A relevance score (0 to 1).

    The function sorts the citations by descending score and displays each with:
        - A numbered header including a normalized relevance percentage.
        - A clickable link to the source.
        - The context text.
        - A horizontal rule divider.
    """
    # Display a main header for the sources section
    st.markdown("## Sources", unsafe_allow_html=True)

    # Sort the citations by score (highest relevance first)
    sorted_citations = sorted(citations, key=lambda c: c.get("score", 0), reverse=False)

    # Loop over each citation and display the formatted content
    for i, citation in enumerate(sorted_citations, start=1):
        # Get raw distance from ChromaDB
        distance = citation.get("score", 0)

        # Convert ChromaDB's cosine distance to a normalized similarity score (0-1)
        similarity_score = 1 - (distance / 2)  # Ensures range 0 (worst) to 1 (best)
        
        # Ensure score is within valid bounds (0-1) before converting to percentage
        normalized_score = round(max(0, min(similarity_score, 1)) * 100, 2)

        # Create a header with the source number and relevance percentage
        st.markdown(
            f"### Source {i} (Relevance: {normalized_score}%)", unsafe_allow_html=True
        )

        # If a URL is present, create a clickable link
        url = citation.get("url", "")
        if url:
            st.markdown(f"[Link to Source]({url})", unsafe_allow_html=True)

        # Display the citation context text
        context_text = citation.get("context", "")
        st.markdown(context_text, unsafe_allow_html=True)

        # Add a horizontal rule to separate sources
        st.markdown("---", unsafe_allow_html=True)



# Helper to format a single citation as a professional reference (AMA-like)
def format_citation(citation, idx=None):
    # Try to use available fields for best formatting
    # Accepts dicts with keys: title, url, year, abstract, id, etc.
    title = citation.get("title") or citation.get("context") or ""
    url = citation.get("url") or citation.get("link") or ""
    year = citation.get("year", "")
    # Try to extract PMID/DOI if present
    pmid = citation.get("id", "")
    # Compose a simple professional citation
    citation_str = f"{title}"
    if year:
        citation_str += f". {year}."
    if url:
        citation_str += f" {url}"
    if pmid:
        citation_str += f" PMID: {pmid}."
    # Numbered if idx is provided
    if idx is not None:
        citation_str = f"{idx}. {citation_str}"
    return citation_str.strip()

# Convert Markdown text to a Word document, with optional citations at the end
def markdown_to_word(markdown_text, citations=None):
    doc = Document()
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add the user question at the top of the document if available
    if st.session_state.get("original_question"):
        h_user_q = doc.add_heading("User Question", level=2)
        h_user_q_run = h_user_q.runs[0]
        h_user_q_run.font.name = 'Calibri'
        h_user_q_run.font.size = Pt(16)
        h_user_q_run.bold = True
        h_user_q.paragraph_format.space_after = Pt(6)

        p_user_q = doc.add_paragraph(st.session_state.original_question)
        p_user_q.paragraph_format.space_after = Pt(6)
        # Ensure runs in this paragraph also use the default font (Calibri 11pt)
        for run_item in p_user_q.runs:
            run_item.font.name = 'Calibri'
            run_item.font.size = Pt(11)

    lines = markdown_text.split("\n")
    for line in lines:
        # Handle headings by counting the leading '#' characters.
        if line.startswith("#"):
            heading_level = len(line) - len(line.lstrip('#'))
            text = line[heading_level:].strip()
            h = doc.add_heading(text, level=min(heading_level, 6))
            if h.runs: # Ensure there's a run to style
                h_run = h.runs[0]
                h_run.font.name = 'Calibri'
                if heading_level == 1:
                    h_run.font.size = Pt(18)
                elif heading_level == 2:
                    h_run.font.size = Pt(16)
                elif heading_level == 3:
                    h_run.font.size = Pt(14)
                else:  # Level 4+
                    h_run.font.size = Pt(12)
                h_run.bold = True
            h.paragraph_format.space_after = Pt(6)
        else:
            # Create a paragraph and process inline markdown
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            segments = re.split(r'(\[.*?\]\(.*?\)|\*\*.*?\*\*|\*.*?\*)', line)
            for seg in segments:
                if seg.startswith("[") and seg.endswith(")"):
                    match = re.match(r'\[([^]]+)\]\(([^)]+)\)', seg)
                    if match:
                        link_text = match.group(1)
                        link_url = match.group(2)
                        add_hyperlink(p, link_url, link_text) # Hyperlink text will be Calibri 11pt
                    else:
                        run = p.add_run(seg) # Inherits Calibri 11pt
                elif seg.startswith("**") and seg.endswith("**"):
                    run = p.add_run(seg[2:-2]) # Inherits Calibri 11pt
                    run.bold = True
                elif seg.startswith("*") and seg.endswith("*"):
                    run = p.add_run(seg[1:-1]) # Inherits Calibri 11pt
                    run.italic = True
                else:
                    run = p.add_run(seg) # Inherits Calibri 11pt

    # Add a list of URLs first if citations are provided
    if citations and isinstance(citations, list) and len(citations) > 0:
        unique_urls = sorted(list(set(c.get("url") or c.get("link") for c in citations if c.get("url") or c.get("link"))))
        if unique_urls:
            doc.add_page_break()
            h_urls = doc.add_heading("Source URLs", level=2)
            if h_urls.runs:
                h_urls_run = h_urls.runs[0]
                h_urls_run.font.name = 'Calibri'
                h_urls_run.font.size = Pt(16)
                h_urls_run.bold = True
            h_urls.paragraph_format.space_after = Pt(6)

            for url_item in unique_urls:
                p_url = doc.add_paragraph()
                p_url.paragraph_format.space_after = Pt(3)
                add_hyperlink(p_url, url_item, url_item) # Hyperlink text will be Calibri 11pt

    # Add detailed citations (references) at the end if provided and non-empty
    if citations and isinstance(citations, list) and len(citations) > 0:
        if not unique_urls: # Add page break only if URLs were not listed (unique_urls might be empty)
             doc.add_page_break()
        
        h_refs = doc.add_heading("References", level=2)
        if h_refs.runs:
            h_refs_run = h_refs.runs[0]
            h_refs_run.font.name = 'Calibri'
            h_refs_run.font.size = Pt(16)
            h_refs_run.bold = True
        h_refs.paragraph_format.space_after = Pt(6)

        for idx, citation in enumerate(citations, 1):
            citation_str = format_citation(citation, idx)
            p_ref = doc.add_paragraph(citation_str, style="List Number")
            p_ref.paragraph_format.space_after = Pt(3)
            for run_item in p_ref.runs:
                run_item.font.name = 'Calibri'
                run_item.font.size = Pt(10)
    return doc


# Extract abstract text from PubMed XML data for a given PMID
async def extract_abstract_from_xml(xml_data: str, pmid: str) -> str:
    try:
        root = ET.fromstring(xml_data)
        for article in root.findall(".//PubmedArticle"):
            medline_citation = article.find("MedlineCitation")
            if medline_citation is not None:
                pmid_element = medline_citation.find("PMID")
                if pmid_element is not None and pmid_element.text == pmid:
                    abstract_element = medline_citation.find(".//Abstract")
                    if abstract_element is not None:
                        abstract_texts = []
                        for elem in abstract_element.findall("AbstractText"):
                            label = elem.get("Label")
                            text = ET.tostring(
                                elem, encoding="unicode", method="text"
                            ).strip()
                            if label is not None and label != "":
                                abstract_texts.append(f"{label}: {text}")
                            else:
                                abstract_texts.append(text)
                        return " ".join(abstract_texts).strip()
        return "No abstract available"
    except ET.ParseError:
        print(f"Error parsing XML for PMID {pmid}")
        return "Error extracting abstract"


# Fetch additional PubMed result IDs if needed
async def fetch_additional_results(
    session: aiohttp.ClientSession,
    search_query: str,
    max_results: int,
    current_count: int,
) -> List[str]:
    additional_needed = max_results - current_count
    url = (
        f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?"
        f"db=pubmed&term={search_query}&sort=relevance&retmode=json&retmax={additional_needed}&"
        f"api_key={st.secrets['pubmed_api_key']}"
    )
    try:
        async with session.get(url, timeout=AIOHTTP_TIMEOUT) as response:
            response.raise_for_status()
            data = await response.json()
            if "esearchresult" in data and isinstance(data["esearchresult"], dict):
                return data["esearchresult"].get("idlist", [])
            logger.warning(f"Unexpected data structure in fetch_additional_results: {data}")
            return []
    except (aiohttp.ClientError, asyncio.TimeoutError) as e:
        print(f"Error fetching additional results: {e}")
        return []


# Fetch PubMed article details and abstract XML data
async def fetch_article_details(
    session: aiohttp.ClientSession,
    id: str,
    details_url: str,
    abstracts_url: str,
    semaphore: asyncio.Semaphore,
) -> Tuple[str, Dict, str]:
    async with semaphore:
        try:
            async with session.get(details_url, timeout=AIOHTTP_TIMEOUT) as details_response:
                details_response.raise_for_status()
                details_data = await details_response.json()
            async with session.get(abstracts_url, timeout=AIOHTTP_TIMEOUT) as abstracts_response:
                abstracts_response.raise_for_status()
                abstracts_data = await abstracts_response.text()
            return id, details_data, abstracts_data
        except (aiohttp.ClientError, asyncio.TimeoutError) as e:
            print(f"Error fetching article details for ID {id}: {e}")
            return id, {}, ""


# Retrieve and filter PubMed abstracts based on search terms and relevance
async def pubmed_abstracts(
    search_terms: str,
    search_type: str, # Parameter kept for signature consistency, not used in query
    max_results_param: int,
    initial_years_back_param: int,
    filter_relevance_param: bool,
    relevance_threshold_param: float,
) -> Tuple[List[Dict[str, str]], bool]:

    async def _internal_search_logic(
        session_param: aiohttp.ClientSession,
        search_terms_for_helper: str,
        year_start_for_helper: int,
        year_end_for_helper: int,
        max_results_to_fetch: int,
        filter_relevance_for_helper: bool,
        relevance_threshold_for_helper: float,
        original_question_for_helper: str
    ) -> List[Dict[str, str]]:
        search_query = f"{search_terms_for_helper}+AND+{year_start_for_helper}[PDAT]:{year_end_for_helper}[PDAT]"
        url = (
            f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?"
            f"db=pubmed&term={search_query}&sort=relevance&retmode=json&retmax={max_results_to_fetch}&"
            f"api_key={st.secrets['pubmed_api_key']}"
        )
        try:
            async with session_param.get(url, timeout=AIOHTTP_TIMEOUT) as response:
                response.raise_for_status()
                data = await response.json()
                if not isinstance(data.get("esearchresult"), dict) or "count" not in data["esearchresult"]:
                    st.error("Unexpected response format from PubMed API (esearch)")
                    return []
                ids = data["esearchresult"].get("idlist", [])
                logger.info(f"PubMed esearch for query '{search_query}' found {len(ids)} article IDs.")
                if not ids:
                    # st.write("No results found for this period.") # Avoid st.write in async backend
                    return []

            articles_data = [] # Renamed from 'articles' to avoid confusion
            semaphore = asyncio.Semaphore(10)
            tasks = []
            for id_str in ids:
                details_url = (
                    f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi?"
                    f"db=pubmed&id={id_str}&retmode=json&api_key={st.secrets['pubmed_api_key']}"
                )
                abstracts_url = (
                    f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?"
                    f"db=pubmed&id={id_str}&retmode=xml&rettype=abstract&api_key={st.secrets['pubmed_api_key']}"
                )
                tasks.append(
                    fetch_article_details(
                        session_param, id_str, details_url, abstracts_url, semaphore
                    )
                )
            results = await asyncio.gather(*tasks)

            processed_articles = []
            for id_str, details_data, abstracts_data in results:
                if "result" in details_data and str(id_str) in details_data["result"]:
                    article_detail = details_data["result"][str(id_str)]
                    year = article_detail["pubdate"].split(" ")[0]
                    if year.isdigit():
                        abstract = await extract_abstract_from_xml(abstracts_data, id_str)
                        article_url = f"https://pubmed.ncbi.nlm.nih.gov/{id_str}"
                        if abstract.strip() and abstract != "No abstract available":
                            processed_articles.append(
                                {
                                    "id": id_str,
                                    "title": article_detail["title"],
                                    "year": year,
                                    "abstract": abstract.strip(),
                                    "link": article_url,
                                }
                            )
                        else:
                            logger.warning(f"No valid abstract found for article ID {id_str}")
            
            if not processed_articles:
                return []

            if filter_relevance_for_helper:
                articles_prompt = "\n".join(
                    [f"ID: {article['id']} - Title: {article['title']}" for article in processed_articles]
                )
                messages = [
                    {"role": "system", "content": "You are an assistant evaluating the relevance of articles to a query. For each article provided, return a relevance score between 0.0 and 1.0 as a JSON object mapping the article's ID to its score, For example return only: {'12345': 0.9, '67890': 0.7}"},
                    {"role": "user", "content": f"Query: {original_question_for_helper}\nArticles:\n{articles_prompt}\n\nReturn a JSON object without additional charactoers."},
                ]
                # st.spinner is UI, cannot be used here. Logging is appropriate.
                logger.info("Filtering PubMed articles for question relevance...")
                try:
                    response = create_chat_completion(messages, model="o3-mini")
                    response_content = response.choices[0].message.content.strip()
                    logger.debug(f"Relevance filtering response content: {response_content}")
                    if not response_content:
                        logger.error("Empty response content received from relevance filtering.")
                        relevance_scores = {}
                    else:
                        try:
                            relevance_scores = json.loads(response_content)
                        except json.JSONDecodeError as e:
                            logger.error(f"Error parsing JSON from relevance filtering: {e}")
                            relevance_scores = {}
                    
                    final_filtered_articles = [
                        article for article in processed_articles
                        if float(relevance_scores.get(str(article["id"]), 0)) >= relevance_threshold_for_helper
                    ]
                    logger.info(f"Found {len(final_filtered_articles)} relevant articles after filtering.")
                    return final_filtered_articles
                except Exception as e:
                    logger.error(f"Error during relevance filtering: {e}")
                    return processed_articles # Fallback to unfiltered if error
            else: # No relevance filtering
                return processed_articles

        except (aiohttp.ClientError, asyncio.TimeoutError) as e:
            logger.error(f"Error fetching PubMed articles for query '{search_query}': {e}")
            # print(f"Error fetching PubMed articles: {e}") # Avoid print in async backend
            return []

    # --- Main logic for pubmed_abstracts ---
    final_articles_to_return = []
    used_older_fallback_articles = False
    current_year = datetime.now().year

    async with aiohttp.ClientSession() as session:
        # Initial search
        start_year_initial = current_year - initial_years_back_param
        logger.info(f"Performing initial PubMed search for '{search_terms}' from {start_year_initial} to {current_year} (last {initial_years_back_param} years).")
        initial_articles_found = await _internal_search_logic(
            session, search_terms, start_year_initial, current_year,
            max_results_param, filter_relevance_param, relevance_threshold_param,
            st.session_state.original_question
        )

        if initial_articles_found:
            final_articles_to_return = initial_articles_found
            logger.info(f"Initial PubMed search found {len(final_articles_to_return)} relevant articles.")
        else: # Initial search found nothing
            if initial_years_back_param < 10: # Max fallback depth is 10 years
                logger.info(f"Initial PubMed search (last {initial_years_back_param} years) yielded no relevant articles. Extending search to 10 years.")
                start_year_extended = current_year - 10 # Fixed 10 years
                
                extended_articles_found = await _internal_search_logic(
                    session, search_terms, start_year_extended, current_year,
                    max_results_param, filter_relevance_param, relevance_threshold_param,
                    st.session_state.original_question
                )

                if extended_articles_found:
                    final_articles_to_return = extended_articles_found
                    used_older_fallback_articles = True
                    logger.info(f"Extended PubMed search (last 10 years) found {len(final_articles_to_return)} relevant articles.")
                else:
                    logger.info("Extended PubMed search (last 10 years) also found no relevant articles.")
            else: # Initial search found nothing, and initial_years_back_param was already >= 10
                logger.info(f"Initial PubMed search (last {initial_years_back_param} years) yielded no relevant articles. Not extending further as search already covered {initial_years_back_param} years.")
    
    # Ensure we return only up to max_results_param
    logger.info(f"Total articles to be added to the database: {len(final_articles_to_return[:max_results_param])}")
    return final_articles_to_return[:max_results_param], used_older_fallback_articles


# Real-time internet search using RapidAPI
def realtime_search(query, domains, max, start_year=2020):
    url = "https://real-time-web-search.p.rapidapi.com/search"
    full_query = f"{query} AND ({domains})"
    start_date = f"{start_year}-01-01"
    end_date = datetime.now().strftime("%Y-%m-%d")
    querystring = {"q": full_query, "limit": max, "from": start_date, "to": end_date}
    headers = {
        "X-RapidAPI-Key": st.secrets["X-RapidAPI-Key"],
        "X-RapidAPI-Host": "real-time-web-search.p.rapidapi.com",
    }
    try:
        response = requests.get(url, headers=headers, params=querystring, timeout=REQUESTS_TIMEOUT)
        response.raise_for_status()
        response_data = response.json().get("data", [])
        urls = [item.get("url") for item in response_data]
        snippets = [
            f"**{item.get('title')}**  \n*{item.get('snippet')}*  \n{item.get('url')} <END OF SITE>"
            for item in response_data
        ]
    except requests.exceptions.RequestException as e:
        st.error(f"RapidAPI real-time search failed to respond: {e}")
        return [], []
    return snippets, urls


#########################################
# OpenAI API Helper Functions for Chat Completions
#########################################
# Retrieve a single response asynchronously
async def get_response(messages, model=experts_model):
    async with aiohttp.ClientSession() as session:
        payload = {"model": model, "messages": messages}
        if model != "o3-mini":
            payload["temperature"] = 0.3
        response = await session.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json=payload,
            timeout=AIOHTTP_TIMEOUT,
        )
        # Ensure response is properly handled if it's not JSON or an error occurred
        if response.status != 200:
            error_text = await response.text()
            logger.error(f"OpenAI API error ({response.status}): {error_text}")
            # Depending on desired behavior, could raise an exception or return a specific error structure
            return {"error": {"message": f"API request failed with status {response.status}: {error_text}", "code": response.status}}
        return await response.json()


# Retrieve multiple responses concurrently
async def get_responses(queries):
    tasks = [get_response(query) for query in queries]
    return await asyncio.gather(*tasks)


#########################################
# Text Cleaning and Source Refinement Functions
#########################################
def clean_text(text):
    text = LOWER_UPPER_PATTERN.sub(r"\1 \2", text)
    text = text.replace("-", " ").replace(" .", ".")
    text = re.sub(r"\s{2,}", " ", text)
    return text


def refine_output(data):
    all_sources = ""
    # Sort the citations by score in descending order and take the top 8
    for i, citation in enumerate(
        sorted(data, key=lambda x: x.get("score", 0), reverse=True)[:8], 1
    ):
        normalized_score = round(citation.get("score", 0) * 100, 2)
        all_sources += f"**Source {i} (Relevance: {normalized_score}%)**\n\n"
        if "url" in citation:
            all_sources += f"[Link to source]({citation['url']})\n\n"
        # Clean and truncate the context text
        cleaned_text = clean_text(citation.get("context", ""))
        truncated_text = (
            cleaned_text[:3000] + "..." if len(cleaned_text) > 3000 else cleaned_text
        )
        all_sources += f"{truncated_text}\n\n"
        # Check for the presence of tabular data
        if "Table" in cleaned_text:
            all_sources += "This source contained tabular data.\n\n"
        all_sources += "---\n\n"
    return all_sources


#########################################
# Local Database Path for the EmbedChain App
#########################################
def get_db_path():
    return tempfile.mkdtemp(prefix="db_")


#########################################
# Helper Function to Extract Expert Information from JSON
#########################################
def extract_expert_info(json_input: str) -> Tuple[List[str], List[str], List[str]]:
    experts, domains, expert_questions = [], [], []
    try:
        data = json.loads(json_input)
        if not isinstance(data, dict) or "rephrased_questions" not in data:
            logger.warning("JSON input for extract_expert_info is missing 'rephrased_questions' key or is not a dict.")
            return experts, domains, expert_questions

        rephrased_list = data["rephrased_questions"]
        if not isinstance(rephrased_list, list):
            logger.warning("'rephrased_questions' is not a list in extract_expert_info.")
            return experts, domains, expert_questions

        for item in rephrased_list:
            if isinstance(item, dict) and all(k in item for k in ["expert", "domain", "question"]):
                experts.append(str(item["expert"]))
                domains.append(str(item["domain"]))
                expert_questions.append(str(item["question"]))
            else:
                logger.warning(f"Skipping malformed item in 'rephrased_questions': {item}")
    except json.JSONDecodeError as e:
        logger.error(f"Failed to decode JSON in extract_expert_info: {e}")
    except Exception as e:
        logger.error(f"Unexpected error in extract_expert_info: {e}")
    return experts, domains, expert_questions


#########################################
# Function to Create Chat Completion (with Caching)
#########################################
@st.cache_data
def create_chat_completion(
    messages,
    google=False,
    model="gpt-4.1",
    frequency_penalty=0,
    logit_bias=None,
    logprobs=False,
    top_logprobs=None,
    max_completion_tokens=5000,
    n=1,
    presence_penalty=0,
    response_format=None,
    seed=None,
    stop=None,
    stream=False,
    include_usage=False,
    temperature=1,
    tools=None,
    tool_choice="none",
    user=None,
):
    
    if model == "o3-mini":
        # Insert the developer message at the start of the messages list.
        messages.insert(0, {
            "role": "developer",
            "content": [
                {
                    "type": "text",
                    "text": "Formatting re-enabled"
                }
            ]
        })
    if google:
        client = OpenAI(
            api_key=st.secrets["GOOGLE_API_KEY"],
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
        )
        params = {
            "model": model,
            "messages": messages,
            "stream": stream,
            "temperature": temperature,
        }
    else:
        client = OpenAI()
        params = {
            "model": model,
            "messages": messages,
            "frequency_penalty": frequency_penalty,
            "logit_bias": logit_bias,
            "logprobs": logprobs,
            "top_logprobs": top_logprobs,
            "max_completion_tokens": max_completion_tokens,
            "n": n,
            "presence_penalty": presence_penalty,
            "response_format": response_format,
            "seed": seed,
            "stop": stop,
            "stream": stream,
            "temperature": temperature,
            "user": user,
        }
    if model == "o3-mini":
        params.pop("temperature", None)
        params["reasoning_effort"] = "medium"

    if stream:
        params["stream_options"] = {"include_usage": include_usage}
    else:
        params.pop("stream_options", None)
    if tools:
        params["tools"] = [{"type": "function", "function": tool} for tool in tools]
        params["tool_choice"] = tool_choice
    if response_format == "json_object":
        params["response_format"] = {"type": "json_object"}
    elif response_format == "text":
        params["response_format"] = {"type": "text"}
    params = {k: v for k, v in params.items() if v is not None}
    completion = client.chat.completions.create(**params)
    return completion


#########################################
# Password Checking for Secure Access
#########################################
def check_password() -> bool:
    if st.secrets["docker"] == "docker":
        st.session_state.password_correct = True
        return True
    if "password" not in st.session_state:
        st.session_state.password = ""
    if "password_correct" not in st.session_state:
        st.session_state.password_correct = False
    if "login_attempts" not in st.session_state:
        st.session_state.login_attempts = 0

    def password_entered() -> None:
        if st.session_state["password"] == st.secrets["password"]:
            st.session_state["password_correct"] = True
            st.session_state.login_attempts = 0
            app = App()
            app.reset()
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False
            st.session_state.login_attempts += 1

    if not st.session_state["password_correct"]:
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        if st.session_state.login_attempts > 0:
            st.error(
                f"ðŸ˜• Password incorrect. Attempts: {st.session_state.login_attempts}"
            )
        st.write(
            "*Please contact David Liebovitz, MD if you need an updated password for access.*"
        )
        return False
    return True


#########################################
# Main Function: Orchestrates the App UI and Workflow
#########################################
def main():
    st.markdown("""
        <style>
            /* Base font and size */
            html, body, [class*="st-"] {
                font-family: 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
                font-size: 1.05rem; /* Slightly larger base font */
            }

            /* Main title style */
            h1 {
                color: #2c3e50; /* A deep, professional blue */
            }

            /* Button styling */
            .stButton>button {
                border-radius: 8px;
                padding: 10px 20px;
                border: 1px solid #2c3e50;
                background-color: #3498db; /* A friendly blue */
                color: white;
                transition: background-color 0.3s ease, box-shadow 0.3s ease;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }
            .stButton>button:hover {
                background-color: #2980b9; /* Darker blue on hover */
                box-shadow: 0 4px 8px rgba(0,0,0,0.15);
            }
            .stButton>button:active {
                background-color: #1f6a9c; /* Even darker blue when active */
            }

            /* Styling for expanders to make them a bit more distinct */
            .stExpander {
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                margin-bottom: 1rem;
            }
            .stExpander header {
                background-color: #f9f9f9;
                border-radius: 8px 8px 0 0;
            }
        </style>
    """, unsafe_allow_html=True)

    st.title("ðŸ” Helpful Answers with AI! âœ¨")
    current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    db_path = get_db_path()
    # query_config = BaseLlmConfig(number_documents=15, model=rag_model)
    # Configure the EmbedChain app based on the selected RAG model
    rag_model = "o3-mini"
    rag_provider = "openai"
    if embedder_provider == "google":
        embedder_config = {"model": embedder_model}
    else:
        embedder_config = {"api_key": embedder_api_key, "model": embedder_model}
    if rag_model == "o3-mini":
        config = {
            "llm": {
                "provider": rag_provider,
                "config": {"model": rag_model, "stream": False, "api_key": rag_key},
            },
            "vectordb": {
                "provider": "chroma",
                "config": {
                    "collection_name": "ai-helper",
                    "dir": db_path,
                    "allow_reset": True,
                },
            },
            "embedder": {
                "provider": embedder_provider,
                "config": embedder_config
            },
            "chunker": {
                "chunk_size": 4500,
                "chunk_overlap": 100,
                "length_function": "len",
                "min_chunk_size": 2000,
            },
        }
    else:
        config = {
            "llm": {
                "provider": rag_provider,
                "config": {
                    "model": rag_model,
                    "temperature": 0.5,
                    "stream": False,
                    "api_key": rag_key,
                },
            },
            "vectordb": {
                "provider": "chroma",
                "config": {
                    "collection_name": "ai-helper",
                    "dir": db_path,
                    "allow_reset": True,
                },
            },
            "embedder": {
                "provider": embedder_provider,
                "config": embedder_config
            },
            "chunker": {
                "chunk_size": 5000,
                "chunk_overlap": 100,
                "length_function": "len",
                "min_chunk_size": 2000,
            },
        }
    # Create a custom initialization that bypasses the alembic migration
    try:
        # First attempt without any special handling
        app = None
        try:
            app = App.from_config(config=config)
        except KeyError as ke:
            if str(ke) == "'script'":
                # Handle the specific alembic error by creating a simpler app
                # Create a basic app without the full config to avoid alembic
                app = App()
                # Then manually set the config components we need
                if "vectordb" in config:
                    app.db = config["vectordb"]
                if "embedder" in config:
                    app.embedder = config["embedder"]
                if "llm" in config and config["llm"].get("provider"):
                    app.llm_provider = config["llm"]["provider"]
                    app.llm_config = config["llm"].get("config", {})
            else:
                st.error(f"Error initializing App: {ke}")
    except Exception as e:
        st.error(f"Error initializing App: {e}")
        app = None
    with st.expander("About this app"):
        st.info(
            """This app interprets a user query and retrieves content from selected internet domains (including PubMed if applicable) for an initial answer and then asks AI personas their opinions on the topic after providing them with updated content. Approaches shown to improve outputs like chain of thought, expert rephrasing, and chain of verification are applied to improve the quality of the responses and to reduce hallucination. Web sites are identified, processed and content selectively retrieved for answers using Real-Time Web Search and the EmbedChain library. App author is David Liebovitz, MD"""
        )
    st.info("Please validate all guidance using the sources!")
    col1, col2 = st.columns([1, 1])
    if check_password():
        # Get user query
        with col1:
            original_query = st.text_area(
                "Ask a nice question...",
                placeholder="Enter your question here...",
                help="Ask any knowledge-based question.",
            )
        st.session_state.original_question = original_query
        find_experts_messages = [
            {"role": "system", "content": system_prompt_expert_questions},
            {"role": "user", "content": original_query},
        ]
        determine_domain_messages = [
            {"role": "system", "content": choose_domain},
            {"role": "user", "content": original_query},
        ]
        first_view = False
        col2.write(" ")
        col2.write(" ")
        col2.write(" ")

        if col2.button("ðŸš€ Start Researching!"):
            # Reset session variables for a new research session
            first_view = True
            st.session_state.articles = []
            st.session_state.pubmed_search_terms = ""
            st.session_state.chosen_domain = ""
            st.session_state.followup_messages = []
            st.session_state.expert_answers = []
            st.session_state.messages1 = []
            st.session_state.messages2 = []
            st.session_state.messages3 = []
            st.session_state.initial_response_thread = []
            st.session_state.final_thread = []
            st.session_state.thread_with_tavily_context = []
            st.session_state.tavily_initial_response = []
            with col1:
                
                if short_use_case == "Helpful PubMed Query":
                    pubmed_messages = [
                        {"role": "system", "content": pubmed_prompt},
                        {"role": "user", "content": original_query},
                    ]
                    response_pubmed_search_terms = create_chat_completion(
                        pubmed_messages, temperature=0.3
                    )
                    pubmed_search_terms = response_pubmed_search_terms.choices[
                        0
                    ].message.content
                    st.session_state.pubmed_search_terms = pubmed_search_terms
                    st.page_link(
                        "https://pubmed.ncbi.nlm.nih.gov/?term="
                        + st.session_state.pubmed_search_terms,
                        label=":green[Click here to open your consensus focused PubMed search!]",
                        icon="ðŸ“š",
                    )                    
                    st.stop()
                    
                if short_use_case == "Helpful Internet Sites":
                    with st.spinner("Determining the best domain for your question..."):
                        restrict_domains_response = create_chat_completion(
                            determine_domain_messages,
                            model=topic_model,
                            temperature=0.3,
                        )
                        st.session_state.chosen_domain = (
                            restrict_domains_response.choices[0].message.content
                        )
                    if (
                        st.session_state.chosen_domain == "medical"
                        and internet_search_provider == "Google"
                    ):
                        domains = (
                            edited_medical_domains
                            if edited_medical_domains != medical_domains
                            else medical_domains
                        )
                    else:
                        if internet_search_provider == "Google":
                            domains = st.session_state.chosen_domain
                    if app is not None:
                        try:
                            if len(app.get_data_sources()) > 0:
                                app.reset()
                        except Exception as e:
                            st.error(f"Error resetting app: {e}; just proceed")

                    search_messages = [
                        {
                            "role": "system",
                            "content": optimize_search_terms_system_prompt,
                        },
                        {
                            "role": "user",
                            "content": f"considering it is {current_datetime}, {original_query}",
                        },
                    ]
                    with st.spinner("Optimizing search terms..."):
                        try:
                            response_google_search_terms = create_chat_completion(
                                search_messages, temperature=0.3
                            )
                        except Exception as e:
                            st.error(f"Error during OpenAI call: {e}")
                    google_search_terms = response_google_search_terms.choices[0].message.content
                    with st.spinner(f'Searching for "{google_search_terms}"...'):
                        if internet_search_provider == "Google":
                            st.session_state.snippets, st.session_state.urls = (
                                realtime_search(
                                    google_search_terms, domains, site_number
                                )
                            )
                        else:
                            three_years_ago = datetime.now() - timedelta(
                                days=3 * 365.25
                            )
                            date_cutoff = three_years_ago.strftime("%Y-%m-%d")
                            search_response = exa.search_and_contents(
                                google_search_terms,
                                text={
                                    "include_html_tags": False,
                                    "max_characters": 1000,
                                },
                                highlights={
                                    "highlights_per_url": 2,
                                    "num_sentences": 5,
                                    "query": "This is the highlight query:",
                                },
                                start_published_date=date_cutoff,
                            )
                            st.session_state.snippets = [
                                result.text for result in search_response.results
                            ]
                            st.session_state.urls = [
                                result.url for result in search_response.results
                            ]
                    with st.expander("Internet Sources", expanded=True):
                        if internet_search_provider == "Google":
                            for snippet in st.session_state.snippets:
                                st.markdown(snippet.replace("<END OF SITE>", ""))
                        else:
                            for i, snippet in enumerate(st.session_state.snippets):
                                st.markdown(
                                    f"### Source {i + 1}: {st.session_state.urls[i]}"
                                )
                                st.markdown(snippet)
                    
                    
                    
                    
                    
                    st.stop()
                    
                # Deeper Dive: Includes PubMed search and web search
                if deeper_dive:
                    with st.spinner("Determining the best domain for your question..."):
                        restrict_domains_response = create_chat_completion(
                            determine_domain_messages,
                            model=topic_model,
                            temperature=0.3,
                        )
                        st.session_state.chosen_domain = (
                            restrict_domains_response.choices[0].message.content
                        )
                    if (
                        st.session_state.chosen_domain == "medical"
                        and internet_search_provider == "Google"
                    ):
                        domains = (
                            edited_medical_domains
                            if edited_medical_domains != medical_domains
                            else medical_domains
                        )
                    else:
                        if internet_search_provider == "Google":
                            domains = st.session_state.chosen_domain
                    try:
                        if len(app.get_data_sources()) > 0:
                            app.reset()
                    except:
                        st.error("Error resetting app; just proceed")

                    search_messages = [
                        {
                            "role": "system",
                            "content": optimize_search_terms_system_prompt,
                        },
                        {
                            "role": "user",
                            "content": f"considering it is {current_datetime}, {original_query}",
                        },
                    ]
                    with st.spinner("Optimizing search terms..."):
                        try:
                            response_google_search_terms = create_chat_completion(
                                search_messages, temperature=0.3
                            )
                        except Exception as e:
                            st.error(f"Error during OpenAI call: {e}")
                    google_search_terms = response_google_search_terms.choices[
                        0
                    ].message.content
                    st.session_state.chosen_domain = (
                        st.session_state.chosen_domain.replace('"', "").replace("'", "")
                    )
                    if st.session_state.chosen_domain == "medical" and short_use_case == "Answer the Question":
                        pubmed_messages = [
                            {"role": "system", "content": pubmed_prompt},
                            {"role": "user", "content": original_query},
                        ]
                        response_pubmed_search_terms = create_chat_completion(
                            pubmed_messages, temperature=0.3
                        )
                        pubmed_search_terms = response_pubmed_search_terms.choices[
                            0
                        ].message.content
                        st.session_state.pubmed_search_terms = pubmed_search_terms
                        with st.spinner(
                            f'Searching PubMed for "{pubmed_search_terms}"...'
                        ):
                            articles_list, older_fetched_flag = asyncio.run(
                                pubmed_abstracts(
                                    search_terms=pubmed_search_terms,
                                    search_type=search_type,
                                    max_results_param=max_results,
                                    initial_years_back_param=years_back,
                                    filter_relevance_param=filter_relevance,
                                    relevance_threshold_param=relevance_threshold,
                                )
                            )
                        st.session_state.articles = articles_list
                        st.session_state.older_pubmed_articles_alert = older_fetched_flag
                        
                        if st.session_state.older_pubmed_articles_alert and st.session_state.articles:
                            st.warning(
                                "Note: The relevant PubMed references identified are primarily older (extending up to 10 years back). "
                                "Please use the PubMed search link to ensure no more recent articles exist."
                            )

                        with st.spinner(
                            "Adding PubMed abstracts to the knowledge base..."
                        ):
                            if st.session_state.articles:
                                for article in st.session_state.articles:
                                    retries = 3
                                    success = False
                                    while retries > 0 and not success:
                                        try:
                                            if not isinstance(article, dict):
                                                raise ValueError(
                                                    "Article is not a valid dictionary."
                                                )
                                            link = article.get("link")
                                            if not link:
                                                raise ValueError(
                                                    "Article does not contain a 'link' key."
                                                )
                                            if app is not None:
                                                app.add(link, data_type="web_page")
                                            success = True
                                        except ValueError as ve:
                                            logger.error(f"Value error: {ve}")
                                            retries -= 1
                                        # Removed the problematic general except Exception here
                                        except Exception as e:
                                            retries -= 1
                                            logger.error(
                                                f"Error adding article {article}: {str(e)}"
                                            )
                                            if retries > 0:
                                                time.sleep(1)
                                            else:
                                                st.error(
                                                    "PubMed results did not meet relevance and recency check. Click the PubMed link to view."
                                                )
                        if not st.session_state.articles:
                            st.warning(
                                "No recent and relevant PubMed articles identified for the knowledge base."
                            )
                            st.page_link(
                                "https://pubmed.ncbi.nlm.nih.gov/?term="
                                + st.session_state.pubmed_search_terms,
                                label="Click here to try directly in PubMed",
                                icon="ðŸ“š",
                            )
                            with st.popover("PubMed Search Terms"):
                                st.write(
                                    f"**Search Strategy:** {st.session_state.pubmed_search_terms}"
                                )
                        else:
                            with st.spinner("Optimizing display of abstracts..."):
                                with st.expander("View PubMed Results While Waiting"):
                                    pubmed_link = (
                                        "https://pubmed.ncbi.nlm.nih.gov/?term="
                                        + st.session_state.pubmed_search_terms
                                    )
                                    st.page_link(
                                        pubmed_link,
                                        label="Click here to view in PubMed",
                                        icon="ðŸ“š",
                                    )
                                    with st.popover("PubMed Search Terms"):
                                        st.write(
                                            f"**Search Strategy:** {st.session_state.pubmed_search_terms}"
                                        )
                                    if st.session_state.older_pubmed_articles_alert and st.session_state.articles: # Duplicating warning here for visibility
                                        st.warning(
                                            "Note: The relevant PubMed references identified are primarily older (extending up to 10 years back). "
                                            "Please use the PubMed search link to ensure no more recent articles exist."
                                        )
                                    for article in st.session_state.articles: # Use st.session_state.articles
                                        st.markdown(
                                            f"### [{article['title']}]({article['link']})"
                                        )
                                        st.write(f"Year: {article['year']}")
                                        st.write(
                                            article["abstract"]
                                            if article["abstract"]
                                            else "No abstract available"
                                        )
                    with st.spinner(f'Searching for "{google_search_terms}"...'):
                        if internet_search_provider == "Google":
                            st.session_state.snippets, st.session_state.urls = (
                                realtime_search(
                                    google_search_terms, domains, site_number
                                )
                            )
                        else:
                            three_years_ago = datetime.now() - timedelta(
                                days=3 * 365.25
                            )
                            date_cutoff = three_years_ago.strftime("%Y-%m-%d")
                            search_response = exa.search_and_contents(
                                google_search_terms,
                                text={
                                    "include_html_tags": False,
                                    "max_characters": 1000,
                                },
                                highlights={
                                    "highlights_per_url": 2,
                                    "num_sentences": 5,
                                    "query": "This is the highlight query:",
                                },
                                start_published_date=date_cutoff,
                            )
                            st.session_state.snippets = [
                                result.text for result in search_response.results
                            ]
                            st.session_state.urls = [
                                result.url for result in search_response.results
                            ]
                    with st.expander("View Reliable Internet Results While Waiting"):
                        for snippet in st.session_state.snippets:
                            st.markdown(snippet.replace("<END OF SITE>", ""))
                    blocked_sites = []
                    with st.spinner("Retrieving full content from web pages..."):
                        for site in st.session_state.urls:
                            if app is not None:
                                try:
                                    app.add(site, data_type="web_page")
                                except Exception:
                                    blocked_sites.append(site)
                    # query_config = BaseLlmConfig(number_documents=15, model=rag_model)
                    with st.spinner("Analyzing retrieved content..."):
                        try:
                            current_datetime = datetime.now().strftime(
                                "%Y-%m-%d %H:%M:%S"
                            )
                            prepare_rag_query_messages = [
                                {"role": "system", "content": prepare_rag_query},
                                {
                                    "role": "user",
                                    "content": f"User query to refine: {original_query}",
                                },
                            ]
                            query_for_rag = create_chat_completion(
                                prepare_rag_query_messages,
                                model="gpt-4o-mini",
                                temperature=0.3,
                            )
                            updated_rag_query = query_for_rag.choices[0].message.content

                        except Exception as e:
                            st.error(f"Error during rag prep {e}")
                        try:
                            if app is not None:
                                citations = app.search(updated_rag_query, where=None, num_documents=20)
                            else:
                                citations = []
                                st.warning("Search functionality is limited due to initialization error.")
                            # with st.expander("View full Citations"):
                            #     display_citations(citations)
                            #     st.write(f'Just the raw{citations}')
                            # print(citations)
                            # citations = [item['metadata']['url'] for item in semantic_results]
                            # filtered_citations = [
                            #     {
                            #         "context": citation.get("context", ""),
                            #         "url": citation.get("metadata", {}).get("url", ""),
                            #         "score": citation.get("metadata", {}).get(
                            #             "score", 0
                            #         ),
                            #     }
                            #     for citation in citations
                            # ]
                            
                            filtered = filter_citations(citations)
                            # with st.expander("View Filtered Citations"):
                            #     display_citations(filtered)
                            #     st.write(f'Just the filtered full{filtered}')
                            st.session_state.citations = filtered
                            # st.markdown(f"**Citations just after filtering:** {filtered_citations}")
                        except Exception as e:
                            st.error(f"Error during semantic query: {e}")
                        try:
                            updated_answer_prompt = rag_prompt2.format(
                                question=original_query,
                                context=st.session_state.citations,
                            )
                            prepare_updated_answer_messages = [
                                {"role": "user", "content": updated_answer_prompt}
                            ]
                            if second_provider == "openai":
                                updated_answer = create_chat_completion(
                                    prepare_updated_answer_messages,
                                    model=second_model,
                                    temperature=0.3,
                                )
                                updated_answer_text = updated_answer.choices[
                                    0
                                ].message.content
                            elif second_provider == "anthropic":
                                client = anthropic.Anthropic(api_key=api_key_anthropic)
                                updated_answer = client.messages.create(
                                    model=second_model,
                                    messages=prepare_updated_answer_messages,
                                    temperature=0.3,
                                    max_tokens=1500,
                                )
                                updated_answer_text = updated_answer.content[0].text
                            elif second_provider == "google":
                                updated_answer = create_chat_completion(
                                    prepare_updated_answer_messages,
                                    google=True,
                                    model=second_model,
                                    temperature=0.3,
                                )
                                updated_answer_text = updated_answer.choices[
                                    0
                                ].message.content
                        except Exception as e:
                            st.error(f"Error during second pass: {e}")
                        # full_response = ""

                        # if citations:
                        #     # st.write(st.session_state.citations)
                        if updated_answer is not None:
                            st.session_state.full_initial_response = updated_answer_text
                    #     first_view = True

                    #     full_response += "\n\n**Sources**:\n"
                    #     sources = []
                    #     for citation in st.session_state.citations:
                    #         # Directly get the URL from the citation dictionary.
                    #         source = citation.get("url", "")
                    #         pattern = re.compile(r"([^/]+)\.[^\.]+\.pdf$")
                    #         match = pattern.search(source)
                    #         if match:
                    #             source = match.group(1) + ".pdf"
                    #         sources.append(source)

                    #     # Remove duplicates by converting to a set and back to a list.
                    #     sources = list(set(sources))
                    #     for source in sources:
                    #         full_response += f"- {source}\n"

                    #     st.session_state.rag_response = full_response

                    # st.session_state.source_chunks = refine_output(citations)
                    container1 = st.container()
                    with container1:
                        st.info(f"Response as of **{current_datetime}:**\n\n")
                        st.markdown(st.session_state.full_initial_response)
                        display_url_list(st.session_state.citations)
                        with st.expander("View Sources"):
                            display_citations(st.session_state.citations)
                        # with st.expander("View Source Excerpts"):
                        #     st.markdown(f'Rag Response: {st.session_state.rag_response}')
                        #     st.markdown(f'Source Chunks: {st.session_state.source_chunks}')
                        #     st.markdown(f'Citations: {st.session_state.citations}')
                        #     st.markdown(st.session_state.source_chunks)
                        if st.button("Create Word Document"):
                            doc = markdown_to_word(
                                st.session_state.rag_response,
                                citations=st.session_state.citations if st.session_state.citations else None,
                            )
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word Document",
                                data=buffer,
                                file_name="prelim_response.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )

                else:
                    # Quick search without extensive PubMed search (Tavily search)
                    with st.spinner("Determining the best domain for your question..."):
                        restrict_domains_response = create_chat_completion(
                            determine_domain_messages,
                            model=topic_model,
                            temperature=0.3,
                        )
                        st.session_state.chosen_domain = (
                            restrict_domains_response.choices[0].message.content
                        )
                    tavily_initial_search_domains = (
                        tavily_domains
                        if st.session_state.chosen_domain == "medical"
                        else ""
                    )
                    try:
                        tavily_client = TavilyClient(
                            api_key=st.secrets["TAVILY_API_KEY"]
                        )
                    except Exception as e:
                        st.error(f"Error during Tavily client initialization: {e}")
                        return
                    with st.spinner(
                        "Accessing reliable internet domains for updates..."
                    ):
                        try:
                            updated_initial_tavily_query = create_chat_completion(
                                [
                                    {
                                        "role": "system",
                                        "content": "Optimize the user question for submission to an online search engine. Return only the optimized question for use in a python pipeline.",
                                    },
                                    {"role": "user", "content": original_query},
                                ]
                            )
                            response = tavily_client.search(
                                query=updated_initial_tavily_query.choices[
                                    0
                                ].message.content,
                                include_domains=tavily_initial_search_domains,
                                search_depth="advanced",
                            )
                        except json.JSONDecodeError as e:
                            print(f"Error decoding JSON: {e}")
                        # st.write(response)
                        st.session_state.tavily_urls = extract_and_format_urls(response)
                        results = response.get("results", [])
                        result_texts = [
                            f"**Title:** {result['title']}\n\n**URL:** {result['url']}\n\n**Content:** {result['content']}\n\n**Relevancy Score:** {result['score']:.2f}\n\n"
                            for result in results
                        ]
                        st.session_state.tavily_initial_response = "\n".join(
                            result_texts
                        )
                    # with st.expander("Retrieved Search Content"):
                    #     st.write(f'Updated Query: {updated_initial_tavily_query.choices[0].message.content}')
                    #     st.write("\n\n")
                    #     st.write(st.session_state.tavily_initial_response)
                    updated_initial_question_with_tavily = (
                        original_query
                        + f" If appropriate, incorporate these updated findings or references from reliable sites; no disclaimers. Users are physicians and expect clear yet dense communication: {st.session_state.tavily_initial_response}"
                    )
                    st.session_state.initial_response_thread.append(
                        {
                            "role": "user",
                            "content": updated_initial_question_with_tavily,
                        }
                    )
                    try:
                        updated_answer = create_chat_completion(
                            st.session_state.initial_response_thread, model="o3-mini"
                        )
                    except Exception as e:
                        st.error(f"Error during second pass: {e}")
                    if updated_answer:
                        current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        full_response = (
                            f"As of **{current_datetime}**:\n\n"
                            + updated_answer.choices[0].message.content
                        )
                        st.session_state.full_initial_response = full_response
                        first_view = True
                container1 = col1.container()
                with container1:
                    if st.session_state.tavily_initial_response:
                        with st.expander("View Tavily Initial Search Results"):
                            st.write(st.session_state.tavily_initial_response)
                        st.info("Initial Response")
                        st.markdown(st.session_state.full_initial_response)
                        st.markdown(st.session_state.tavily_urls)
                        if st.button("Create Word Document for Quick Search"):
                            doc = markdown_to_word(
                                st.session_state.full_initial_response,
                                citations=st.session_state.citations if st.session_state.citations else None,
                            )
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word Document for Quick Search",
                                data=buffer,
                                file_name="full_initial_response.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
        with col2:
            if st.session_state.rag_response or st.session_state.full_initial_response:
                # RAGAS Model Settings
                
                hallucination_check=st.button("Validate Response Section 1 against Sources", help="The Faithfulness Score = Number of claims supported by the sources / Total number of claims in the response")
                if hallucination_check:
                    #### Ragas Scoring

                    pattern = r"^.*2\. Additional Insights from the Model's Knowledge.*$"
                    # Split at the first line MATCHING the plain text phrase (robust to markdown)
                    split_content = re.split(pattern, st.session_state.full_initial_response, maxsplit=1, flags=re.MULTILINE)
                    section1 = split_content[0].rstrip()
                    sample = SingleTurnSample(
                        response=f'User question: {st.session_state.original_question} Response: {section1}',
                        reference=str(st.session_state.citations),
                        )
                    
                    sample_faithfulness = SingleTurnSample(
                        user_input=st.session_state.original_question,
                        response=section1,
                        retrieved_contexts=[str(st.session_state.citations)],
                    )

                    rubrics = {
                        "score1_description": "There is no hallucination in the response. The response is fully supported by the reference.",
                        "score2_description": "Factual statements are supported by the reference but the response is not fully accurate and lacks important details.",
                        "score3_description": "There are some factual statements that are not present in the reference.",
                        "score4_description": "The response contains some factual errors and lacks important details based on the reference.",
                        "score5_description": "The model adds new information and statements that contradict the reference.",
                    }
                    scorer = RubricsScore(rubrics=rubrics, llm=evaluator_llm)
                    scorer_faithfulness = Faithfulness(llm=evaluator_llm)
                    # await scorer.single_turn_ascore(sample)
                    # metric = AspectCritic(name="summary_accuracy",llm=evaluator_llm, definition="Identify the Best Answer from Retrieved Context section. Assess if the sources support the response in this section.")
                    # test_data = SingleTurnSample(**test_data)
                    # with st.spinner("Evaluating response..."):
                    #     pass_or_fail = metric.single_turn_ascore(test_data)
                    #     st.write(f'**RAGAS Score:** {pass_or_fail}')
                        
                    async def evaluate_ragas_metrics(): # Renamed function
                        # scorer is RubricsScore, scorer_faithfulness is Faithfulness
                        # Call internal _single_turn_ascore to get full result objects
                        # Pass empty list for callbacks as it's expected by the internal methods
                        rubric_result_obj = await scorer._single_turn_ascore(sample, callbacks=[])
                        faithfulness_result_obj = await scorer_faithfulness._single_turn_ascore(sample_faithfulness, callbacks=[])
                        return rubric_result_obj, faithfulness_result_obj

                    # evaluate_ragas_metrics returns the direct scores
                    direct_rubric_score, direct_faithfulness_score = asyncio.run(evaluate_ragas_metrics())
                    
                    current_rubric_score = int(direct_rubric_score) 
                    current_faithfulness_score = float(direct_faithfulness_score)

                    # Display existing summary messages based on scores
                    if current_rubric_score == 1:
                        st.success("Section 1 is supported by the sources (Rubric Score: 1).")
                    elif current_rubric_score == 2:
                        st.error("Caution: Factual statements supported, but Section 1 may lack accuracy/details (Rubric Score: 2). Confirm with references.")
                    elif current_rubric_score == 3:
                        st.warning("Caution: Some factual statements in Section 1 may not be fully supported (Rubric Score: 3). Confirm with references.")
                    elif current_rubric_score == 4:
                        st.warning("Warning: Section 1 may contain factual errors/lack details (Rubric Score: 4). Confirm with references.")
                    elif current_rubric_score == 5:
                        st.error("Warning!!! Section 1 may add new information contradicting sources (Rubric Score: 5). Confirm with references.")
                    else:
                        st.error(f"Error: Unable to evaluate the response based on rubrics (Rubric Score: {current_rubric_score}).")
                    
                    if current_faithfulness_score > 0.9:
                        st.success(f"**Faithfulness Score:** {current_faithfulness_score:.3f} (High confidence in factual consistency with sources).")
                    else:
                        st.warning(f"**Faithfulness Score:** {current_faithfulness_score:.3f}. Review carefully, some assertions might not be fully backed by provided sources.")

                    # New expander for detailed RAGAS results
                    with st.expander("View RAGAS Evaluation Details"):
                        st.subheader("Rubric Score Details")
                        st.markdown(f"**Overall Rubric Score:** {current_rubric_score}")
                        # Derive reason from the rubrics dictionary
                        rubric_reason_key = f"score{current_rubric_score}_description"
                        rubric_reason = rubrics.get(rubric_reason_key, "Specific reason not found for this score.")
                        st.markdown("**Reasoning (from predefined rubrics):**")
                        st.markdown(rubric_reason)
                        
                        st.divider()
                        
                        st.subheader("Faithfulness Score Details")
                        st.markdown(f"**Overall Faithfulness Score:** {current_faithfulness_score:.3f}")
                        st.markdown("The Faithfulness Score measures the factual consistency of the generated answer against the provided context. A higher score indicates better alignment.")
                        
                        # Add more detailed explanation of how faithfulness is calculated
                        st.markdown("### How Faithfulness is Calculated")
                        st.markdown("""
                        1. **Statement Generation**: The response is broken down into individual factual statements.
                        2. **Statement Verification**: Each statement is checked against the retrieved context.
                        3. **Verdict Assignment**: Each statement receives a verdict (1 if supported by context, 0 if not).
                        4. **Score Calculation**: Final score = Number of supported statements / Total number of statements.
                        
                        A score closer to 1.0 means most statements in the response are directly supported by the provided sources.
                        A lower score indicates the response contains statements not found in or contradicted by the sources.
                        """)
                        
                        # Add statement breakdown and verdicts
                        st.markdown("### Statement Breakdown")
                        
                        # Generate example statements with verdicts for demonstration
                        # In a real implementation, these would come from the RAGAS evaluation
                        with st.spinner("Analyzing statements..."):
                            statements_prompt = [
                                {"role": "system", "content": "You are an AI assistant that breaks down text into individual factual statements. For the given text, extract 5-8 key factual claims as a JSON list of strings. Each statement should be self-contained and represent a single factual claim. Return a JSON object with a 'statements' array. If you can't find any statements, include at least 3 general claims from the text."},
                                {"role": "user", "content": f"Extract factual statements from this text:\n\n{section1}"}
                            ]
                                
                            try:
                                client = OpenAI()
                                statements_response = client.chat.completions.create(
                                    model="gpt-4o-mini",
                                    messages=statements_prompt,
                                    response_format={"type": "json_object"},
                                    temperature=0.1
                                )
                                statements_json = statements_response.choices[0].message.content
                                    
                                try:
                                    parsed_json = json.loads(statements_json)
                                    statements = parsed_json.get("statements", [])
                                        
                                    # Fallback if no statements were extracted
                                    if not statements:
                                        st.warning("No specific factual statements were identified. Using general statements from the text.")
                                        # Create some basic statements from the text
                                        statements = [
                                            f"The text discusses {st.session_state.original_question}",
                                            "The response provides information based on available sources",
                                            "The text contains factual content related to the query"
                                        ]
                                except json.JSONDecodeError:
                                    st.warning("Could not parse JSON response. Using general statements.")
                                    statements = [
                                        f"The text discusses {st.session_state.original_question}",
                                        "The response provides information based on available sources",
                                        "The text contains factual content related to the query"
                                    ]
                                
                                # For each statement, determine if it's supported by the context
                                verdicts = []
                                for statement in statements:
                                    try:
                                        verdict_prompt = [
                                            {"role": "system", "content": "You are an AI assistant that determines if a statement is supported by the provided context. Return a JSON object with keys 'verdict' (1 if supported, 0 if not) and 'reason' (brief explanation)."},
                                            {"role": "user", "content": f"Statement: {statement}\n\nContext: {str(st.session_state.citations)}\n\nIs this statement supported by the context? Return only a JSON object."}
                                        ]
                                        
                                        client = OpenAI()
                                        verdict_response = client.chat.completions.create(
                                            model="gpt-4o-mini",
                                            messages=verdict_prompt,
                                            response_format={"type": "json_object"},
                                            temperature=0.1
                                        )
                                        
                                        verdict_json = verdict_response.choices[0].message.content
                                        
                                        try:
                                            verdict_data = json.loads(verdict_json)
                                            verdict = verdict_data.get("verdict", 0)
                                            # Ensure verdict is either 0 or 1
                                            if not isinstance(verdict, int) or verdict not in [0, 1]:
                                                if isinstance(verdict, str) and verdict.lower() in ["true", "yes", "supported"]:
                                                    verdict = 1
                                                elif isinstance(verdict, str) and verdict.lower() in ["false", "no", "not supported"]:
                                                    verdict = 0
                                                else:
                                                    verdict = 0
                                                    
                                            verdicts.append({
                                                "statement": statement,
                                                "verdict": verdict,
                                                "reason": verdict_data.get("reason", "No reason provided")
                                            })
                                        except json.JSONDecodeError:
                                            st.warning(f"Could not parse verdict JSON for statement: {statement[:50]}...")
                                            verdicts.append({
                                                "statement": statement,
                                                "verdict": 0,
                                                "reason": "Error processing verdict"
                                            })
                                    except Exception as e:
                                        st.warning(f"Error evaluating statement: {str(e)}")
                                        verdicts.append({
                                            "statement": statement,
                                            "verdict": 0,
                                            "reason": f"Error during evaluation: {str(e)}"
                                        })
                                
                                # Display statements and verdicts in a flat list (no nested expanders)
                                if verdicts:
                                    st.markdown("The following statements were evaluated:")
                                    
                                    # Create a table for all statements
                                    statement_data = []
                                    for i, v in enumerate(verdicts, 1):
                                        verdict_icon = "âœ…" if v["verdict"] == 1 else "âŒ"
                                        statement_data.append({
                                            "Statement": f"{verdict_icon} {v['statement']}",
                                            "Verdict": "Supported" if v["verdict"] == 1 else "Not supported",
                                            "Reason": v["reason"]
                                        })
                                    
                                    # Display as a DataFrame
                                    import pandas as pd
                                    df = pd.DataFrame(statement_data)
                                    st.dataframe(df, use_container_width=True)
                                    
                                    # Calculate and display the faithfulness score based on these verdicts
                                    supported = sum(1 for v in verdicts if v["verdict"] == 1)
                                    total = len(verdicts)
                                    calculated_score = supported / total if total > 0 else 0
                                    st.markdown(f"**Calculated score:** {calculated_score:.3f} ({supported} supported statements out of {total} total)")
                                    
                                    # Compare with RAGAS score
                                    st.markdown(f"**RAGAS faithfulness score:** {current_faithfulness_score:.3f}")
                                    if abs(calculated_score - current_faithfulness_score) > 0.2:
                                        st.info("Note: There's a significant difference between our calculated score and the RAGAS score. This could be due to differences in statement extraction or evaluation methods.")
                                else:
                                    st.error("No statements were successfully evaluated. Please try again.")
                                
                            except Exception as e:
                                st.error(f"Error analyzing statements: {str(e)}")
                                st.markdown("Unable to display statement breakdown. Please try again later.")
                        
                        # Add explanation about limitations
                        st.markdown("### Limitations")
                        st.markdown("""
                        - The metric evaluates factual consistency, not completeness or relevance.
                        - Complex or nuanced statements may be difficult to verify automatically.
                        - The evaluation depends on the quality of the statement breakdown process.
                        - Statements that are common knowledge but not in the sources may be marked as unfaithful.
                        - The statement extraction and verification process shown here is a demonstration and may differ from RAGAS's internal implementation.
                        """)
                
                #     st.session_state.ragas_score = current_rubric_score # If you still need this elsewhere
                # if st.session_state.full_initial_response:
                #     st.sidebar.write(f'**Score (1 or 0) if Section 1 is faithful to Sources:** {st.session_state.ragas_score}')
                initial_followup = st.checkbox(
                    "Ask Follow-Up Questions for Initial Response"
                )
                if initial_followup:
                    add_internet_content = True
                    # prelim_response = st.session_state.full_initial_response
                    formatted_output = []
                    for citation in st.session_state.citations:
                        try:
                            source_metadata, source_text = citation
                            metadata_details = (
                                "\n".join(
                                    f"{key.capitalize()}: {value}"
                                    for key, value in source_metadata.items()
                                )
                                if isinstance(source_metadata, dict)
                                else f"Metadata: {source_metadata}"
                            )
                            formatted_output.append(
                                f"{metadata_details}\nSource text: {source_text}\n---"
                            )
                        except Exception as e:
                            formatted_output.append(
                                f"Error processing citation: {citation}\nError: {str(e)}\n---"
                            )
                    formatted_output_str = "\n".join(formatted_output)
                    prelim_followup_prompt2 = prelim_followup_prompt.format(
                        prior_question=original_query,
                        evidence=formatted_output_str,
                        prior_answer=st.session_state.full_initial_response,
                    )
                    if not st.session_state.initial_response_thread:
                        st.session_state.initial_response_thread.append(
                            {"role": "system", "content": prelim_followup_prompt2}
                        )
                    if initial_followup_question := st.chat_input("Ask followup!"):
                        if add_internet_content:
                            try:
                                tavily_client = TavilyClient(
                                    api_key=st.secrets["TAVILY_API_KEY"]
                                )
                            except Exception as e:
                                st.error(
                                    f"Error during Tavily client initialization: {e}"
                                )
                                return
                            with st.spinner(
                                "Retrieving additional internet content..."
                            ):
                                try:
                                    updated_tavily_query = create_chat_completion(
                                        [
                                            {
                                                "role": "system",
                                                "content": "Combine user inputs (an initial and then followup question) into one question optimized for searching online. Return only the optimized question which will be used in a python pipeline.",
                                            },
                                            {
                                                "role": "user",
                                                "content": f"{original_query} and {initial_followup_question}",
                                            },
                                        ]
                                    )
                                    response = tavily_client.search(
                                        query=updated_tavily_query.choices[
                                            0
                                        ].message.content,
                                        include_domains=tavily_domains,
                                        search_depth="advanced",
                                    )
                                except json.JSONDecodeError as e:
                                    print(f"Error decoding JSON: {e}")
                                results = response.get("results", [])
                                result_texts = [
                                    f"**Title:** {result['title']}\n\n**URL:** {result['url']}\n\n**Content:** {result['content']}\n\n**Relevancy Score:** {result['score']:.2f}\n\n"
                                    for result in results
                                ]
                                st.session_state.tavily_followup_response = "\n".join(
                                    result_texts
                                )
                            with st.expander("New Retrieved Search Content"):
                                st.write(
                                    f"Updated Query: {updated_tavily_query.choices[0].message.content}"
                                )
                                st.write("\n\n")
                                st.write(st.session_state.tavily_followup_response)
                            updated_followup_question = (
                                initial_followup_question
                                + f"Here's more of what I found online but wnat your thoughts: {st.session_state.tavily_followup_response}"
                            )
                            st.session_state.thread_with_tavily_context = (
                                st.session_state.initial_response_thread
                            )
                            st.session_state.thread_with_tavily_context.append(
                                {"role": "user", "content": updated_followup_question}
                            )
                            initial_followup_messages = (
                                st.session_state.thread_with_tavily_context
                            )
                        else:
                            st.session_state.initial_response_thread.append(
                                {"role": "user", "content": initial_followup_question}
                            )
                            initial_followup_messages = (
                                st.session_state.initial_response_thread
                            )
                        with st.chat_message("user"):
                            st.markdown(initial_followup_question)
                        with st.chat_message("assistant"):
                            client = OpenAI()
                            try:
                                stream = client.chat.completions.create(
                                    model="o3-mini",
                                    messages=[
                                        {"role": m["role"], "content": m["content"]}
                                        for m in initial_followup_messages
                                    ],
                                    stream=True,
                                )
                                response = st.write_stream(stream)
                            except Exception as e:
                                st.error(f"Error during OpenAI call: {e}")
                                return
                            st.session_state.initial_response_thread.append(
                                {"role": "assistant", "content": response}
                            )
                    with st.expander("View full follow-up thread"):
                        for message in st.session_state.initial_response_thread:
                            if message["role"] != "system":
                                emoji = role_emojis.get(message["role"], "â“")
                                st.write(
                                    f"{emoji} {message['role'].capitalize()}: {message['content']}"
                                )
                    if st.session_state.initial_response_thread:
                        if st.checkbox("Download Followup Conversation"):
                            full_followup_conversation = ""
                            for message in st.session_state.initial_response_thread:
                                if message["role"] != "system":
                                    emoji = role_emojis.get(message["role"], "â“")
                                    full_followup_conversation += f"{emoji} {message['role'].capitalize()}: {message['content']}\n\n"
                            html = markdown2.markdown(
                                full_followup_conversation, extras=["tables"]
                            )
                            st.download_button(
                                "Download Followup Conversation",
                                html,
                                "followup_conversation.html",
                                "text/html",
                            )
                if not initial_followup:
                    if st.button("Ask 3 AI Expert Personas for Opinions"):
                        prelim_response = st.session_state.full_initial_response + str(
                            st.session_state.citations
                        )
                        try:
                            completion = create_chat_completion(
                                messages=find_experts_messages,
                                model=experts_model,
                                temperature=0.3,
                                response_format="json_object",
                            )
                        except Exception as e:
                            st.error(f"Error during OpenAI call: {e}")
                            return
                        json_output = completion.choices[0].message.content
                        experts, domains, expert_questions = extract_expert_info(
                            json_output
                        )
                        st.session_state.experts = experts
                        updated_expert1_system_prompt = expert1_system_prompt.format(
                            expert=experts[0], domain=domains[0]
                        )
                        updated_expert2_system_prompt = expert2_system_prompt.format(
                            expert=experts[1], domain=domains[1]
                        )
                        updated_expert3_system_prompt = expert3_system_prompt.format(
                            expert=experts[2], domain=domains[2]
                        )
                        updated_question1 = expert_questions[0]
                        updated_question2 = expert_questions[1]
                        updated_question3 = expert_questions[2]
                        prelim_response = (
                            st.session_state.rag_response
                            + st.session_state.source_chunks
                        )
                        expert1_messages = [
                            {
                                "role": "system",
                                "content": updated_expert1_system_prompt,
                            },
                            {
                                "role": "user",
                                "content": updated_question1
                                + "Here's what I already found online: "
                                + prelim_response,
                            },
                        ]
                        st.session_state.messages1 = expert1_messages
                        expert2_messages = [
                            {
                                "role": "system",
                                "content": updated_expert2_system_prompt,
                            },
                            {
                                "role": "user",
                                "content": updated_question2
                                + "Here's what I already found online: "
                                + prelim_response,
                            },
                        ]
                        st.session_state.messages2 = expert2_messages
                        expert3_messages = [
                            {
                                "role": "system",
                                "content": updated_expert3_system_prompt,
                            },
                            {
                                "role": "user",
                                "content": updated_question3
                                + "Here's what I already found online: "
                                + prelim_response,
                            },
                        ]
                        st.session_state.messages3 = expert3_messages
                        with st.spinner("Waiting for experts to respond..."):
                            st.session_state.expert_answers = asyncio.run(
                                get_responses(
                                    [
                                        expert1_messages,
                                        expert2_messages,
                                        expert3_messages,
                                    ]
                                )
                            )
        with col1:
            if not first_view:
                try:
                    if (
                        st.session_state.pubmed_search_terms
                        and st.session_state.articles
                    ):
                        with st.expander("View PubMed Results Added to Knowledge Base"):
                            pubmed_link = (
                                "https://pubmed.ncbi.nlm.nih.gov/?term="
                                + st.session_state.pubmed_search_terms
                            )
                            st.page_link(
                                pubmed_link,
                                label="Click here to view in PubMed",
                                icon="ðŸ“š",
                            )
                            with st.popover("PubMed Search Terms"):
                                st.write(
                                    f"**Search Strategy:** {st.session_state.pubmed_search_terms}"
                                )
                            if st.session_state.older_pubmed_articles_alert and st.session_state.articles:
                                st.warning(
                                    "Note: The relevant PubMed references identified are primarily older (extending up to 10 years back). "
                                    "Please use the PubMed search link to ensure no more recent articles exist."
                                )
                            for article in st.session_state.articles:
                                st.markdown(
                                    f"### [{article['title']}]({article['link']})"
                                )
                                st.write(f"Year: {article['year']}")
                                st.write(
                                    article["abstract"]
                                    if article["abstract"]
                                    else "No abstract available"
                                )
                except:
                    st.write(
                        "No Relevant PubMed articles to display - if topic works, API may be down!"
                    )
                    pubmed_link = (
                        "https://pubmed.ncbi.nlm.nih.gov/?term="
                        + st.session_state.pubmed_search_terms
                    )
                    st.page_link(
                        pubmed_link, label="Click here to try in PubMed", icon="ðŸ“š"
                    )
                    with st.popover("PubMed Search Terms"):
                        st.write(
                            f"**Search Strategy:** {st.session_state.pubmed_search_terms}"
                        )
                with st.expander("View Internet Results Added to Knowledge Base"):
                    for snippet in st.session_state.snippets:
                        st.markdown(snippet.replace("<END OF SITE>", ""))
                if st.session_state.full_initial_response:
                    # st.info(f"Response as of **{current_datetime}:**\n\n")
                    container1 = st.container()
                    with container1:
                        st.info(f"Response as of **{current_datetime}:**\n\n")
                        st.markdown(st.session_state.full_initial_response)
                        if st.session_state.citations:
                            display_url_list(st.session_state.citations)
                            with st.expander("View Source Excerpts"):
                                display_citations(st.session_state.citations)
                        if st.session_state.tavily_initial_response:
                            st.markdown(st.session_state.tavily_urls)
                            with st.expander("Retrieved Search Content"):
                                # st.write(f'Updated Query: {updated_initial_tavily_query.choices[0].message.content}')
                                # st.write("\n\n")
                                st.write(st.session_state.tavily_initial_response)
                        if st.button("Create Word File"):
                            doc = markdown_to_word(
                                st.session_state.full_initial_response,
                                citations=st.session_state.citations if st.session_state.citations else None,
                            )
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File",
                                data=buffer,
                                file_name="prelim_response.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        # with st.expander("View Source Excerpts"):
                        #     st.markdown(st.session_state.source_chunks)
        with col2:
            if st.session_state.expert_answers:
                container2 = st.container()
                with container2:
                    st.info("AI Expert Persona Responses")
                    with st.expander(f"AI {st.session_state.experts[0]} Perspective"):
                        expert_0 = st.session_state.expert_answers[0]["choices"][0][
                            "message"
                        ]["content"]
                        st.write(expert_0)
                        st.session_state.messages1.append(
                            {"role": "assistant", "content": expert_0}
                        )
                        if st.button("Create Word File for AI Expert 1"):
                            doc = markdown_to_word(
                                expert_0,
                                citations=st.session_state.citations if st.session_state.citations else None,
                            )
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File for AI Expert 1",
                                data=buffer,
                                file_name="AI_expert_1.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                    with st.expander(f"AI {st.session_state.experts[1]} Perspective"):
                        expert_1 = st.session_state.expert_answers[1]["choices"][0][
                            "message"
                        ]["content"]
                        st.write(expert_1)
                        st.session_state.messages2.append(
                            {"role": "assistant", "content": expert_1}
                        )
                        if st.button("Create Word File for AI Expert 2"):
                            doc = markdown_to_word(
                                expert_1,
                                citations=st.session_state.citations if st.session_state.citations else None,
                            )
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File for AI Expert 2",
                                data=buffer,
                                file_name="AI_expert_2.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                    with st.expander(f"AI {st.session_state.experts[2]} Perspective"):
                        expert_2 = st.session_state.expert_answers[2]["choices"][0][
                            "message"
                        ]["content"]
                        st.write(expert_2)
                        st.session_state.messages3.append(
                            {"role": "assistant", "content": expert_2}
                        )
                        if st.button("Create Word File for AI Expert 3"):
                            doc = markdown_to_word(
                                expert_2,
                                citations=st.session_state.citations if st.session_state.citations else None,
                            )
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File for AI Expert 3",
                                data=buffer,
                                file_name="AI_expert_3.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )


#########################################
# Main Program Entry Point
#########################################
if __name__ == "__main__":
    main()
