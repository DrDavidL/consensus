
#########################################
# Import Libraries and Setup
#########################################
import asyncio
import json
import re
import tempfile
import time
from datetime import datetime, timedelta

from docx import Document
from io import BytesIO

import aiohttp
import requests
import streamlit as st
import anthropic
from openai import OpenAI
from exa_py import Exa
import markdown2
import xml.etree.ElementTree as ET
from typing import List, Tuple, Dict
from tavily import TavilyClient

from embedchain import App
from embedchain.config import BaseLlmConfig

import logging
from requests.exceptions import RequestException

#########################################
# Global Variables and Logging Configuration
#########################################
role_emojis = {
    "user": "👤",
    "assistant": "🤖",
}

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

#########################################
# Import Prompts for AI Guidance and Search
#########################################
from prompts import (
    system_prompt_expert_questions,
    expert1_system_prompt,
    expert2_system_prompt,
    expert3_system_prompt,
    optimize_search_terms_system_prompt,
    optimize_pubmed_search_terms_system_prompt,
    cutting_edge_pubmed_prompt,
    prepare_rag_query,
    rag_prompt2,
    choose_domain,
    medical_domains,
    prelim_followup_prompt,
    tavily_domains,
    evaluate_response_prompt
)

#########################################
# Streamlit App Configuration and API Keys Setup
#########################################
st.set_page_config(
    page_title='Helpful AI',
    layout='wide',
    page_icon=':stethoscope:',
    initial_sidebar_state='collapsed'
)
api_key = st.secrets["OPENAI_API_KEY"]
api_key_anthropic = st.secrets["ANTHROPIC_API_KEY"]
exa = Exa(st.secrets["EXA_API_KEY"])

#########################################
# Session State Initialization
#########################################
if "snippets" not in st.session_state:
    st.session_state["snippets"] = []
if "urls" not in st.session_state:
    st.session_state["urls"] = []
if "rag_response" not in st.session_state:
    st.session_state["rag_response"] = ""
if "source_chunks" not in st.session_state:
    st.session_state.source_chunks = ''
if "experts" not in st.session_state:
    st.session_state.experts = []
if "messages1" not in st.session_state:
    st.session_state.messages1 = []
if "messages2" not in st.session_state:
    st.session_state.messages2 = []
if "messages3" not in st.session_state:
    st.session_state.messages3 = []
if "followup_messages" not in st.session_state:
    st.session_state.followup_messages = []
if "final_thread" not in st.session_state:
    st.session_state.final_thread = []
if "expert_number" not in st.session_state:
    st.session_state.expert_number = 0
if "expert_answers" not in st.session_state:
    st.session_state.expert_answers = []
if "original_question" not in st.session_state:
    st.session_state.original_question = ""
if "chosen_domain" not in st.session_state:
    st.session_state.chosen_domain = ""
if "articles" not in st.session_state:
    st.session_state.articles = []
if "pubmed_search_terms" not in st.session_state:
    st.session_state.pubmed_search_terms = ""
if "full_initial_response" not in st.session_state:
    st.session_state.full_initial_response = ""
if "initial_response_thread" not in st.session_state:
    st.session_state.initial_response_thread = []
if "citations" not in st.session_state:
    st.session_state.citations = []
if "thread_with_tavily_context" not in st.session_state:
    st.session_state.thread_with_tavily_context = []
if "tavily_followup_response" not in st.session_state:
    st.session_state.tavily_followup_response = ""
if "tavily_initial_response" not in st.session_state:
    st.session_state.tavily_initial_response = ""

#########################################
# Sidebar Configuration: UI Elements & Settings
#########################################
with st.sidebar:
    # Toggle for subject area model
    topic_model_choice = st.toggle(
        "Subject Area: Use GPT-4o",
        help="Toggle to use GPT-4o model for determining if medical; otherwise, 4o-mini."
    )
    if topic_model_choice:
        st.write("GPT-4o model selected.")
        topic_model = "gpt-4o"
    else:
        st.write("GPT-4o-mini model selected.")
        topic_model = "gpt-4o-mini"
    st.divider()

    # Web search settings
    search_type = "all"
    with st.sidebar.popover("Web Search Settings"):
        site_number = st.number_input(
            "Number of web pages to retrieve:",
            min_value=1, max_value=20, value=10, step=1
        )
        internet_search_provider = st.radio(
            "Internet search provider:",
            options=["Google", "Exa"],
            horizontal=True,
            help="Only specific Google domains are used for retrieving current Medical or General Knowledge. Exa.ai is a new type of search tool that predicts relevant sites; domain filtering not yet added here."
        )
        if internet_search_provider == "Google":
            st.info("Web domains used for medical questions.")
            edited_medical_domains = st.text_area(
                "Edit domains (maintain format pattern):",
                medical_domains, height=200
            )
    st.divider()

    # PubMed search settings
    st.sidebar.info("PubMed Search Settings")
    years_back = st.slider(
        "Years Back for PubMed Search",
        min_value=1, max_value=10, value=4, step=1,
        help="Set the number of years back to search PubMed."
    )
    st.divider()
    max_results = st.slider(
        "Number of Abstracts to Review",
        min_value=3, max_value=20, value=10, step=1,
        help="Set the number of abstracts to review."
    )
    st.divider()
    filter_relevance = st.toggle(
        "Filter Relevance of PubMed searching",
        value=True, help="Toggle to deselect."
    )
    if filter_relevance:
        relevance_threshold = st.slider(
            "Relevance Threshold",
            min_value=0.3, max_value=1.0, value=0.8, step=0.05,
            help="Set the minimum relevance score to consider an item relevant."
        )
    else:
        relevance_threshold = 0.75
        st.write("Top sources will be added to the database regardless.")
    st.divider()

    # Technical settings for embedder model
    st.sidebar.info("More Technical Settings")
    st.divider()
    embedder_model_choice = st.toggle(
        "Embedder Model: Use text-embedding-3-large",
        help="Toggle to use text-embedding-3-large."
    )
    if embedder_model_choice:
        st.write("text-embedding-3-large model selected.")
        embedder_model = "text-embedding-3-large"
    else:
        st.write("text-embedding-3-small model selected.")
        embedder_model = "text-embedding-3-small"
    st.divider()
    st.info("GPT-4o-mini performs well for other options. For more complex synthesis, stay with GPT-4o or use Claude-3.5 Sonnet.")
    
    # RAG model options
    rag_model_choice = st.radio(
        "RAG Model Options",
        ["GPT-4o-mini", "GPT-4o", "Gemini-1.5"],
        index=1,
        help="Select the RAG model to use for the AI responses."
    )
    if rag_model_choice == "GPT-4o":
        st.write("GPT-4o model selected.")
        rag_model = "gpt-4o"
        rag_provider = "openai"
        rag_key = api_key
    elif rag_model_choice == "Gemini-1.5":
        st.write("Gemini-1.5 model selected.")
        rag_model = "gemini-1.5-pro-latest"
        rag_provider = "google"
        rag_key = st.secrets["GOOGLE_API_KEY"]
    elif rag_model_choice == "GPT-4o-mini":
        st.write("GPT-4o-mini model selected.")
        rag_model = "gpt-4o-mini"
        rag_provider = "openai"
        rag_key = api_key
    st.divider()

    # Second review model options
    second_review_model = st.radio(
        "Second Review Model Options",
        ["GPT-4o-mini", "GPT-4o", "o3-mini", "Claude-3.5 Sonnet", "Gemini-1.5"],
        index=2,
        help="Select the RAG model to use for the AI responses."
    )
    if second_review_model == "GPT-4o":
        st.write("GPT-4o model selected.")
        second_model = "gpt-4o"
        second_provider = "openai"
        second_key = api_key
    elif second_review_model == "Claude-3.5 Sonnet":
        st.write("Claude-3-5-sonnet-latest model selected.")
        second_model = "claude-3-5-sonnet-latest"
        second_provider = "anthropic"
        second_key = api_key_anthropic
    elif second_review_model == "GPT-4o-mini":
        st.write("GPT-4o-mini model selected.")
        second_model = "gpt-4o-mini"
        second_provider = "openai"
        second_key = api_key
    elif second_review_model == "Gemini-1.5":
        st.write("Gemini-1.5 model selected.")
        second_model = "gemini-1.5-pro-latest"
        second_provider = "google"
        second_key = st.secrets["GOOGLE_API_KEY"]
    elif second_review_model == "o3-mini":
        st.write("o3-mini reasoning model selected.")
        second_model = "o3-mini"
        second_provider = "openai"
        second_key = api_key
    st.divider()

    # Expert personas model choice
    experts_model_choice = st.toggle(
        "3 AI Experts Model: Use GPT-4o",
        help="Toggle to use GPT-4o model for expert responses; otherwise, o3-mini with reasoning."
    )
    if experts_model_choice:
        st.write("GPT-4o model selected.")
        experts_model = "gpt-4o"
    else:
        st.write("o3-mini reasoning model selected.")
        experts_model = "o3-mini"

#########################################
# Utility Functions
#########################################

# Convert Markdown text to a Word document
def markdown_to_word(markdown_text):
    doc = Document()
    lines = markdown_text.split("\n")
    for line in lines:
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    return doc

# Extract abstract text from PubMed XML data for a given PMID
async def extract_abstract_from_xml(xml_data: str, pmid: str) -> str:
    try:
        root = ET.fromstring(xml_data)
        for article in root.findall(".//PubmedArticle"):
            medline_citation = article.find("MedlineCitation")
            if medline_citation:
                pmid_element = medline_citation.find("PMID")
                if pmid_element is not None and pmid_element.text == pmid:
                    abstract_element = medline_citation.find(".//Abstract")
                    if abstract_element is not None:
                        abstract_texts = []
                        for elem in abstract_element.findall("AbstractText"):
                            label = elem.get("Label")
                            text = ET.tostring(elem, encoding='unicode', method='text').strip()
                            if label:
                                abstract_texts.append(f"{label}: {text}")
                            else:
                                abstract_texts.append(text)
                        return " ".join(abstract_texts).strip()
        return "No abstract available"
    except ET.ParseError:
        print(f"Error parsing XML for PMID {pmid}")
        return "Error extracting abstract"

# Fetch additional PubMed result IDs if needed
async def fetch_additional_results(session: aiohttp.ClientSession, search_query: str, max_results: int, current_count: int) -> List[str]:
    additional_needed = max_results - current_count
    url = (
        f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?"
        f"db=pubmed&term={search_query}&sort=relevance&retmode=json&retmax={additional_needed}&"
        f"api_key={st.secrets['pubmed_api_key']}"
    )
    try:
        async with session.get(url) as response:
            response.raise_for_status()
            data = await response.json()
            return data['esearchresult'].get('idlist', [])
    except (aiohttp.ClientError, asyncio.TimeoutError) as e:
        print(f"Error fetching additional results: {e}")
        return []

# Fetch PubMed article details and abstract XML data
async def fetch_article_details(session: aiohttp.ClientSession, id: str, details_url: str, abstracts_url: str, semaphore: asyncio.Semaphore) -> Tuple[str, Dict, str]:
    async with semaphore:
        try:
            async with session.get(details_url) as details_response:
                details_response.raise_for_status()
                details_data = await details_response.json()
            async with session.get(abstracts_url) as abstracts_response:
                abstracts_response.raise_for_status()
                abstracts_data = await abstracts_response.text()
            return id, details_data, abstracts_data
        except (aiohttp.ClientError, asyncio.TimeoutError) as e:
            print(f"Error fetching article details for ID {id}: {e}")
            return id, {}, ""

# Retrieve and filter PubMed abstracts based on search terms and relevance
async def pubmed_abstracts(
    search_terms: str,
    search_type: str = "all",
    max_results: int = max_results,
    years_back: int = years_back,
    filter_relevance: bool = filter_relevance,
    relevance_threshold: float = relevance_threshold
) -> List[Dict[str, str]]:
    current_year = datetime.now().year
    start_year = current_year - years_back
    search_query = f"{search_terms}+AND+{start_year}[PDAT]:{current_year}[PDAT]"
    url = (
        f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?"
        f"db=pubmed&term={search_query}&sort=relevance&retmode=json&retmax={max_results}&"
        f"api_key={st.secrets['pubmed_api_key']}"
    )
    async with aiohttp.ClientSession() as session:
        try:
            async with session.get(url) as response:
                response.raise_for_status()
                data = await response.json()
                if 'esearchresult' not in data or 'count' not in data['esearchresult']:
                    st.error("Unexpected response format from PubMed API")
                    return []
                ids = data['esearchresult'].get('idlist', [])
                if not ids:
                    st.write("No results found.")
                    return []
            articles = []
            semaphore = asyncio.Semaphore(5)
            tasks = []
            for id in ids:
                details_url = (
                    f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi?"
                    f"db=pubmed&id={id}&retmode=json&api_key={st.secrets['pubmed_api_key']}"
                )
                abstracts_url = (
                    f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?"
                    f"db=pubmed&id={id}&retmode=xml&rettype=abstract&api_key={st.secrets['pubmed_api_key']}"
                )
                tasks.append(fetch_article_details(session, id, details_url, abstracts_url, semaphore))
            results = await asyncio.gather(*tasks)
            filtered_articles = []
            for id, details_data, abstracts_data in results:
                if 'result' in details_data and str(id) in details_data['result']:
                    article = details_data['result'][str(id)]
                    year = article['pubdate'].split(" ")[0]
                    if year.isdigit():
                        abstract = await extract_abstract_from_xml(abstracts_data, id)
                        article_url = f"https://pubmed.ncbi.nlm.nih.gov/{id}"
                        if abstract.strip() and abstract != "No abstract available":
                            filtered_articles.append({
                                'id': id,
                                'title': article['title'],
                                'year': year,
                                'abstract': abstract.strip(),
                                'link': article_url
                            })
            # Filter articles based on relevance score if required
            if filter_relevance:
                relevant_articles = []
                for article in filtered_articles:
                    messages = [
                        {
                            'role': 'system',
                            'content': "You are an assistant evaluating relevance of abstracts to a query. You only return a score between 0 and 1."
                        },
                        {
                            'role': 'user',
                            'content': f"Query: {st.session_state.original_question}\nAbstract: {article['abstract']}\nNow, respond only with a relevance score between 0 and 1 representing the likelihood the answer is found in this article. Sample response: 0.9"
                        }
                    ]
                    with st.spinner("Filtering PubMed articles for question relevance"):
                        try:
                            response = create_chat_completion(messages, model="gpt-4o-mini", temperature=0.3)
                            relevance_score = float(response.choices[0].message.content.strip())
                            if relevance_score >= relevance_threshold:
                                relevant_articles.append(article)
                        except Exception as e:
                            logger.error(f"Error filtering article: {e}")
                            continue
                articles = [
                    {
                        'id': a['id'],
                        'title': a['title'],
                        'link': a['link'],
                        'year': a['year'],
                        'abstract': a['abstract']
                    } for a in relevant_articles
                ]
            else:
                articles = [
                    {
                        'id': a['id'],
                        'title': a['title'],
                        'link': a['link'],
                        'year': a['year'],
                        'abstract': a['abstract']
                    } for a in filtered_articles
                ]
        except aiohttp.ClientError as e:
            st.error(f"Error connecting to PubMed API: {e}")
            return []
        except Exception as e:
            st.error(f"An unexpected error occurred: {e}")
            return []
    return articles[:max_results]

# Real-time internet search using RapidAPI
def realtime_search(query, domains, max, start_year=2020):
    url = "https://real-time-web-search.p.rapidapi.com/search"
    full_query = f"{query} AND ({domains})"
    start_date = f"{start_year}-01-01"
    end_date = datetime.now().strftime('%Y-%m-%d')
    querystring = {
        "q": full_query,
        "limit": max,
        "from": start_date,
        "to": end_date
    }
    headers = {
        "X-RapidAPI-Key": st.secrets["X-RapidAPI-Key"],
        "X-RapidAPI-Host": "real-time-web-search.p.rapidapi.com",
    }
    try:
        response = requests.get(url, headers=headers, params=querystring)
        response.raise_for_status()
        response_data = response.json().get('data', [])
        urls = [item.get('url') for item in response_data]
        snippets = [
            f"**{item.get('title')}**  \n*{item.get('snippet')}*  \n{item.get('url')} <END OF SITE>"
            for item in response_data
        ]
    except requests.exceptions.RequestException as e:
        st.error(f"RapidAPI real-time search failed to respond: {e}")
        return [], []
    return snippets, urls

#########################################
# OpenAI API Helper Functions for Chat Completions
#########################################
# Retrieve a single response asynchronously
async def get_response(messages, model=experts_model):
    async with aiohttp.ClientSession() as session:
        payload = {'model': model, 'messages': messages}
        if model != "o3-mini":
            payload["temperature"] = 0.3
        response = await session.post(
            'https://api.openai.com/v1/chat/completions',
            headers={
                'Authorization': f'Bearer {api_key}',
                'Content-Type': 'application/json',
            },
            json=payload
        )
        return await response.json()

# Retrieve multiple responses concurrently
async def get_responses(queries):
    tasks = [get_response(query) for query in queries]
    return await asyncio.gather(*tasks)

#########################################
# Text Cleaning and Source Refinement Functions
#########################################
def clean_text(text):
    text = re.sub(r"([a-z])([A-Z])", r"\1 \2", text)
    text = text.replace('-', ' ').replace(' .', '.')
    text = re.sub(r"\s{2,}", " ", text)
    return text

def refine_output(data):
    all_sources = ""
    for i, (text, info) in enumerate(sorted(data, key=lambda x: x[1]['score'], reverse=True)[:8], 1):
        normalized_score = round(info['score'] * 100, 2)
        all_sources += f"**Source {i} (Relevance: {normalized_score}%)**\n\n"
        if 'url' in info:
            all_sources += f"[Link to source]({info['url']})\n\n"
        cleaned_text = clean_text(text)
        truncated_text = cleaned_text[:3000] + "..." if len(cleaned_text) > 3000 else cleaned_text
        all_sources += f"{truncated_text}\n\n"
        if "Table" in cleaned_text:
            all_sources += "This source contained tabular data.\n\n"
        all_sources += "---\n\n"
    return all_sources

#########################################
# Local Database Path for the EmbedChain App
#########################################
def get_db_path():
    return tempfile.mkdtemp(prefix="db_")

#########################################
# Helper Function to Extract Expert Information from JSON
#########################################
def extract_expert_info(json_input):
    data = json.loads(json_input)
    experts = []
    domains = []
    expert_questions = []
    for item in data['rephrased_questions']:
        experts.append(item['expert'])
        domains.append(item['domain'])
        expert_questions.append(item['question'])
    return experts, domains, expert_questions

#########################################
# Function to Create Chat Completion (with Caching)
#########################################
@st.cache_data
def create_chat_completion(
    messages,
    google=False,
    model="gpt-4o",
    frequency_penalty=0,
    logit_bias=None,
    logprobs=False,
    top_logprobs=None,
    max_tokens=None,
    n=1,
    presence_penalty=0,
    response_format=None,
    seed=None,
    stop=None,
    stream=False,
    include_usage=False,
    temperature=1,
    tools=None,
    tool_choice="none",
    user=None
):
    if google:
        client = OpenAI(
            api_key=st.secrets["GOOGLE_API_KEY"],
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
        )
        params = {
            "model": model,
            "messages": messages,
            "stream": stream,
            "temperature": temperature,
        }
    else:
        client = OpenAI()
        params = {
            "model": model,
            "messages": messages,
            "frequency_penalty": frequency_penalty,
            "logit_bias": logit_bias,
            "logprobs": logprobs,
            "top_logprobs": top_logprobs,
            "max_tokens": max_tokens,
            "n": n,
            "presence_penalty": presence_penalty,
            "response_format": response_format,
            "seed": seed,
            "stop": stop,
            "stream": stream,
            "temperature": temperature,
            "user": user
        }
    if model == "o3-mini":
        params.pop("temperature", None)
    if stream:
        params["stream_options"] = {"include_usage": include_usage}
    else:
        params.pop("stream_options", None)
    if tools:
        params["tools"] = [{"type": "function", "function": tool} for tool in tools]
        params["tool_choice"] = tool_choice
    if response_format == "json_object":
        params["response_format"] = {"type": "json_object"}
    elif response_format == "text":
        params["response_format"] = {"type": "text"}
    else:
        params.pop("response_format", None)
    params = {k: v for k, v in params.items() if v is not None}
    completion = client.chat.completions.create(**params)
    return completion

#########################################
# Password Checking for Secure Access
#########################################
def check_password() -> bool:
    if st.secrets["docker"] == "docker":
        st.session_state.password_correct = True
        return True
    if "password" not in st.session_state:
        st.session_state.password = ""
    if "password_correct" not in st.session_state:
        st.session_state.password_correct = False
    if "login_attempts" not in st.session_state:
        st.session_state.login_attempts = 0

    def password_entered() -> None:
        if st.session_state["password"] == st.secrets["password"]:
            st.session_state["password_correct"] = True
            st.session_state.login_attempts = 0
            app = App()
            app.reset()
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False
            st.session_state.login_attempts += 1

    if not st.session_state["password_correct"]:
        st.text_input("Password", type="password", on_change=password_entered, key='password')
        if st.session_state.login_attempts > 0:
            st.error(f"😕 Password incorrect. Attempts: {st.session_state.login_attempts}")
        st.write("*Please contact David Liebovitz, MD if you need an updated password for access.*")
        return False
    return True

#########################################
# Main Function: Orchestrates the App UI and Workflow
#########################################
def main():
    st.title('Helpful Answers with AI!')
    db_path = get_db_path()
    # Configure the EmbedChain app based on the selected RAG model
    if rag_model == "o3-mini":
        config = {
            "llm": {
                "provider": rag_provider,
                "config": {"model": rag_model, "stream": False, "api_key": rag_key},
            },
            "vectordb": {
                "provider": "chroma",
                "config": {"collection_name": "ai-helper", "dir": db_path, "allow_reset": True},
            },
            "embedder": {
                "provider": "openai",
                "config": {"api_key": api_key, "model": embedder_model},
            },
            "chunker": {
                "chunk_size": 1000,
                "chunk_overlap": 50,
                "length_function": "len",
                "min_chunk_size": 200,
            },
        }
    else:
        config = {
            "llm": {
                "provider": rag_provider,
                "config": {"model": rag_model, "temperature": 0.5, "stream": False, "api_key": rag_key},
            },
            "vectordb": {
                "provider": "chroma",
                "config": {"collection_name": "ai-helper", "dir": db_path, "allow_reset": True},
            },
            "embedder": {
                "provider": "openai",
                "config": {"api_key": api_key, "model": embedder_model},
            },
            "chunker": {
                "chunk_size": 1000,
                "chunk_overlap": 50,
                "length_function": "len",
                "min_chunk_size": 200,
            },
        }
    app = App.from_config(config=config)
    with st.expander("About this app"):
        st.info(
            """This app interprets a user query and retrieves content from selected internet domains (including PubMed if applicable) for an initial answer and then asks AI personas their opinions on the topic after providing them with updated content. Approaches shown to improve outputs like chain of thought, expert rephrasing, and chain of verification are applied to improve the quality of the responses and to reduce hallucination. Web sites are identified, processed and content selectively retrieved for answers using Real-Time Web Search and the EmbedChain library. The LLM model is GPT-4o from OpenAI. App author is David Liebovitz, MD"""
        )
    st.info("Please validate all guidance using the sources!")
    col1, col2 = st.columns([1, 1])
    if check_password():
        # Get user query
        with col1:
            original_query = st.text_area(
                'Ask a nice question...',
                placeholder='Enter your question here...',
                help="Ask any knowledge-based question."
            )
        st.session_state.original_question = original_query
        find_experts_messages = [
            {'role': 'system', 'content': system_prompt_expert_questions},
            {'role': 'user', 'content': original_query}
        ]
        determine_domain_messages = [
            {'role': 'system', 'content': choose_domain},
            {'role': 'user', 'content': original_query}
        ]
        first_view = False
        col2.write(" ")
        col2.write(" ")
        col2.write(" ")
        # Check if cutting-edge PubMed research should be included
        if st.sidebar.checkbox(
            "Include Cutting-Edge Research in PubMed (default is consensus review articles)",
            help="Check to include latest, not yet consensus, articles in the search for medical content.",
            value=False
        ):
            pubmed_prompt = cutting_edge_pubmed_prompt
        else:
            pubmed_prompt = optimize_pubmed_search_terms_system_prompt
        deeper_dive = st.sidebar.checkbox("Deeper Dive", help="Check to include more extensive searching.", value=True)
        if col2.button('Begin Research'):
            # Reset session variables for a new research session
            st.session_state.articles = []
            st.session_state.pubmed_search_terms = ""
            st.session_state.chosen_domain = ""
            st.session_state.followup_messages = []
            st.session_state.expert_answers = []
            st.session_state.messages1 = []
            st.session_state.messages2 = []
            st.session_state.messages3 = []
            st.session_state.initial_response_thread = []
            st.session_state.final_thread = []
            st.session_state.thread_with_tavily_context = []
            st.session_state.tavily_initial_response = []
            with col1:
                # Deeper Dive: Includes PubMed search and web search
                if deeper_dive:
                    with st.spinner('Determining the best domain for your question...'):
                        restrict_domains_response = create_chat_completion(
                            determine_domain_messages,
                            model=topic_model,
                            temperature=0.3
                        )
                        st.session_state.chosen_domain = restrict_domains_response.choices[0].message.content
                    if st.session_state.chosen_domain == "medical" and internet_search_provider == "Google":
                        domains = edited_medical_domains if edited_medical_domains != medical_domains else medical_domains
                    else:
                        if internet_search_provider == "Google":
                            domains = st.session_state.chosen_domain
                    try:
                        if len(app.get_data_sources()) > 0:
                            app.reset()
                    except:
                        st.error("Error resetting app; just proceed")
                    current_datetime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    search_messages = [
                        {'role': 'system', 'content': optimize_search_terms_system_prompt},
                        {'role': 'user', 'content': f'considering it is {current_datetime}, {original_query}'}
                    ]
                    with st.spinner('Optimizing search terms...'):
                        try:
                            response_google_search_terms = create_chat_completion(search_messages, temperature=0.3)
                        except Exception as e:
                            st.error(f"Error during OpenAI call: {e}")
                    google_search_terms = response_google_search_terms.choices[0].message.content
                    st.session_state.chosen_domain = st.session_state.chosen_domain.replace('"', '').replace("'", '')
                    if st.session_state.chosen_domain == "medical":
                        pubmed_messages = [
                            {'role': 'system', 'content': pubmed_prompt},
                            {'role': 'user', 'content': original_query}
                        ]
                        response_pubmed_search_terms = create_chat_completion(pubmed_messages, temperature=0.3)
                        pubmed_search_terms = response_pubmed_search_terms.choices[0].message.content
                        st.session_state.pubmed_search_terms = pubmed_search_terms
                        with st.spinner(f'Searching PubMed for "{pubmed_search_terms}"...'):
                            articles = asyncio.run(pubmed_abstracts(pubmed_search_terms, search_type, max_results, years_back))
                        st.session_state.articles = articles
                        with st.spinner("Adding PubMed abstracts to the knowledge base..."):
                            if articles:
                                for article in articles:
                                    retries = 3
                                    success = False
                                    while retries > 0 and not success:
                                        try:
                                            if not isinstance(article, dict):
                                                raise ValueError("Article is not a valid dictionary.")
                                            link = article.get("link")
                                            if not link:
                                                raise ValueError("Article does not contain a 'link' key.")
                                            app.add(link, data_type='web_page')
                                            success = True
                                        except Exception as e:
                                            retries -= 1
                                            logger.error(f"Error adding article {article}: {str(e)}")
                                            if retries > 0:
                                                time.sleep(1)
                                            else:
                                                st.error("PubMed results did not meet relevance and recency check. Click the PubMed link to view.")
                        if not articles:
                            st.warning("No recent and relevant PubMed articles identified for the knowledge base.")
                            st.page_link(
                                "https://pubmed.ncbi.nlm.nih.gov/?term=" + st.session_state.pubmed_search_terms,
                                label="Click here to try directly in PubMed",
                                icon="📚"
                            )
                            with st.popover("PubMed Search Terms"):
                                st.write(f'**Search Strategy:** {st.session_state.pubmed_search_terms}')
                        else:
                            with st.spinner("Optimizing display of abstracts..."):
                                with st.expander("View PubMed Results Added to Knowledge Base"):
                                    pubmed_link = "https://pubmed.ncbi.nlm.nih.gov/?term=" + st.session_state.pubmed_search_terms
                                    st.page_link(pubmed_link, label="Click here to view in PubMed", icon="📚")
                                    with st.popover("PubMed Search Terms"):
                                        st.write(f'**Search Strategy:** {st.session_state.pubmed_search_terms}')
                                    for article in articles:
                                        st.markdown(f"### [{article['title']}]({article['link']})")
                                        st.write(f"Year: {article['year']}")
                                        st.write(article['abstract'] if article['abstract'] else "No abstract available")
                    with st.spinner(f'Searching for "{google_search_terms}"...'):
                        if internet_search_provider == "Google":
                            st.session_state.snippets, st.session_state.urls = realtime_search(google_search_terms, domains, site_number)
                        else:
                            three_years_ago = datetime.now() - timedelta(days=3 * 365.25)
                            date_cutoff = three_years_ago.strftime("%Y-%m-%d")
                            search_response = exa.search_and_contents(
                                google_search_terms,
                                text={"include_html_tags": False, "max_characters": 1000},
                                highlights={"highlights_per_url": 2, "num_sentences": 5, "query": "This is the highlight query:"},
                                start_published_date=date_cutoff
                            )
                            st.session_state.snippets = [result.text for result in search_response.results]
                            st.session_state.urls = [result.url for result in search_response.results]
                    with st.expander("View Internet Results Added to Knowledge Base"):
                        for snippet in st.session_state.snippets:
                            st.markdown(snippet.replace('<END OF SITE>', ''))
                    blocked_sites = []
                    with st.spinner('Retrieving full content from web pages...'):
                        for site in st.session_state.urls:
                            try:
                                app.add(site, data_type='web_page')
                            except Exception as e:
                                blocked_sites.append(site)
                    query_config = BaseLlmConfig(number_documents=15, model=rag_model)
                    with st.spinner('Analyzing retrieved content...'):
                        try:
                            current_datetime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                            prepare_rag_query_messages = [
                                {'role': 'system', 'content': prepare_rag_query},
                                {'role': 'user', 'content': f'User query to refine: {original_query}'}
                            ]
                            query_for_rag = create_chat_completion(prepare_rag_query_messages, model="o3-mini", temperature=0.3)
                            updated_rag_query = query_for_rag.choices[0].message.content 
                        except Exception as e:
                            st.error(f"Error during rag prep {e}")
                        try:
                            answer, citations = app.query(updated_rag_query, config=query_config, citations=True)
                            st.session_state.citations = citations
                        except Exception as e:
                            st.error(f"Error during rag query: {e}")
                        try:
                            updated_answer_prompt = rag_prompt2.format(
                                question=original_query,
                                prelim_answer=answer,
                                context=citations
                            )
                            prepare_updated_answer_messages = [{'role': 'user', 'content': updated_answer_prompt}]
                            if second_provider == "openai":
                                updated_answer = create_chat_completion(prepare_updated_answer_messages, model=second_model, temperature=0.3)
                                updated_answer_text = updated_answer.choices[0].message.content
                            elif second_provider == "anthropic":
                                client = anthropic.Anthropic(api_key=api_key_anthropic)
                                updated_answer = client.messages.create(model=second_model, messages=prepare_updated_answer_messages, temperature=0.3, max_tokens=1500)
                                updated_answer_text = updated_answer.content[0].text
                            elif second_provider == "google":
                                updated_answer = create_chat_completion(prepare_updated_answer_messages, google=True, model=second_model, temperature=0.3)
                                updated_answer_text = updated_answer.choices[0].message.content
                        except Exception as e:
                            st.error(f"Error during second pass: {e}")
                    full_response = ""
                    if answer:
                        full_response = f"From retrieved content, **{current_datetime}:**\n\n{answer} \n\n"
                        if updated_answer is not None:
                            full_response += "\n\n **Second Pass Consolidation**:\n\n*********************\n\n" + updated_answer_text
                            st.session_state.full_initial_response = full_response
                        first_view = True
                    if citations:
                        full_response += "\n\n**Sources**:\n"
                        sources = []
                        for i, citation in enumerate(citations):
                            source = citation[1]["url"]
                            pattern = re.compile(r"([^/]+)\.[^\.]+\.pdf$")
                            match = pattern.search(source)
                            if match:
                                source = match.group(1) + ".pdf"
                            sources.append(source)
                        sources = list(set(sources))
                        for source in sources:
                            full_response += f"- {source}\n"
                        st.session_state.rag_response = full_response
                    st.session_state.source_chunks = refine_output(citations)
                    container1 = st.container()
                    with container1:
                        st.info("Initial Response")
                        st.markdown(st.session_state.rag_response)
                        if st.button("Create Word Document"):
                            doc = markdown_to_word(st.session_state.rag_response)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word Document",
                                data=buffer,
                                file_name="prelim_response.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        with st.expander("View Source Excerpts"):
                            st.markdown(st.session_state.source_chunks)
                else:
                    # Quick search without extensive PubMed search (Tavily search)
                    with st.spinner('Determining the best domain for your question...'):
                        restrict_domains_response = create_chat_completion(determine_domain_messages, model=topic_model, temperature=0.3)
                        st.session_state.chosen_domain = restrict_domains_response.choices[0].message.content
                    tavily_initial_search_domains = tavily_domains if st.session_state.chosen_domain == "medical" else ""
                    try:
                        tavily_client = TavilyClient(api_key=st.secrets["TAVILY_API_KEY"])
                    except Exception as e:
                        st.error(f"Error during Tavily client initialization: {e}")
                        return
                    with st.spinner("Retrieving internet content..."):
                        try:
                            updated_initial_tavily_query = create_chat_completion(
                                [
                                    {'role': 'system', 'content': "Optimize the user question for submission to an online search engine. Return only the optimized question for use in a python pipeline."},
                                    {'role': 'user', 'content': original_query}
                                ]
                            )
                            response = tavily_client.search(
                                query=updated_initial_tavily_query.choices[0].message.content,
                                include_domains=tavily_initial_search_domains,
                                search_depth="advanced"
                            )
                        except json.JSONDecodeError as e:
                            print(f"Error decoding JSON: {e}")
                        results = response.get("results", [])
                        result_texts = [
                            f"**Title:** {result['title']}\n\n**URL:** {result['url']}\n\n**Content:** {result['content']}\n\n**Relevancy Score:** {result['score']:.2f}\n\n"
                            for result in results
                        ]
                        st.session_state.tavily_initial_response = "\n".join(result_texts)
                    with st.expander("Retrieved Search Content"):
                        st.write(f'Updated Query: {updated_initial_tavily_query.choices[0].message.content}')
                        st.write("\n\n")
                        st.write(st.session_state.tavily_initial_response)
                    updated_initial_question_with_tavily = original_query + f"Use the following current internet results: {st.session_state.tavily_initial_response}"
                    st.session_state.initial_response_thread.append({"role": "user", "content": updated_initial_question_with_tavily})
                    try:
                        updated_answer = create_chat_completion(st.session_state.initial_response_thread, model="gpt-4o", temperature=0.3)
                    except Exception as e:
                        st.error(f"Error during second pass: {e}")
                    if updated_answer:
                        current_datetime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                        full_response = f"As of **{current_datetime}**:\n\n" + updated_answer.choices[0].message.content
                        st.session_state.full_initial_response = full_response
                        first_view = True
                container1 = col1.container()
                with container1:
                    if st.session_state.tavily_initial_response:
                        with st.expander("View Tavily Initial Search Results"):
                            st.write(st.session_state.tavily_initial_response)
                        st.info("Initial Response")
                        st.markdown(st.session_state.full_initial_response)
                        if st.button("Create Word Document for Quick Search"):
                            doc = markdown_to_word(st.session_state.full_initial_response)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word Document for Quick Search",
                                data=buffer,
                                file_name="full_initial_response.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
        with col2:
            if st.session_state.rag_response or st.session_state.full_initial_response:
                initial_followup = st.checkbox("Ask Follow-Up Questions for Initial Response")
                if initial_followup:
                    add_internet_content = True
                    prelim_response = st.session_state.rag_response + st.session_state.source_chunks
                    formatted_output = []
                    for citation in st.session_state.citations:
                        try:
                            source_metadata, source_text = citation
                            metadata_details = (
                                "\n".join(f"{key.capitalize()}: {value}" for key, value in source_metadata.items())
                                if isinstance(source_metadata, dict)
                                else f"Metadata: {source_metadata}"
                            )
                            formatted_output.append(f"{metadata_details}\nSource text: {source_text}\n---")
                        except Exception as e:
                            formatted_output.append(f"Error processing citation: {citation}\nError: {str(e)}\n---")
                    formatted_output_str = "\n".join(formatted_output)
                    prelim_followup_prompt2 = prelim_followup_prompt.format(
                        prior_question=original_query,
                        evidence=formatted_output_str,
                        prior_answer=st.session_state.rag_response
                    )
                    if not st.session_state.initial_response_thread:
                        st.session_state.initial_response_thread.append({"role": "system", "content": prelim_followup_prompt2})
                    if initial_followup_question := st.chat_input("Ask followup!"):
                        if add_internet_content:
                            try:
                                tavily_client = TavilyClient(api_key=st.secrets["TAVILY_API_KEY"])
                            except Exception as e:
                                st.error(f"Error during Tavily client initialization: {e}")
                                return
                            with st.spinner("Retrieving additional internet content..."):
                                try:
                                    updated_tavily_query = create_chat_completion(
                                        [
                                            {'role': 'system', 'content': "Combine user inputs (an initial and then followup question) into one question optimized for searching online. Return only the optimized question which will be used in a python pipeline."},
                                            {'role': 'user', 'content': f'{original_query} and {initial_followup_question}'}
                                        ]
                                    )
                                    response = tavily_client.search(
                                        query=updated_tavily_query.choices[0].message.content,
                                        include_domains=tavily_domains,
                                        search_depth="advanced"
                                    )
                                except json.JSONDecodeError as e:
                                    print(f"Error decoding JSON: {e}")
                                results = response.get("results", [])
                                result_texts = [
                                    f"**Title:** {result['title']}\n\n**URL:** {result['url']}\n\n**Content:** {result['content']}\n\n**Relevancy Score:** {result['score']:.2f}\n\n"
                                    for result in results
                                ]
                                st.session_state.tavily_followup_response = "\n".join(result_texts)
                            with st.expander("New Retrieved Search Content"):
                                st.write(f'Updated Query: {updated_tavily_query.choices[0].message.content}')
                                st.write("\n\n")
                                st.write(st.session_state.tavily_followup_response)
                            updated_followup_question = initial_followup_question + f"Here's more of what I found online but wnat your thoughts: {st.session_state.tavily_followup_response}"
                            st.session_state.thread_with_tavily_context = st.session_state.initial_response_thread
                            st.session_state.thread_with_tavily_context.append({"role": "user", "content": updated_followup_question})
                            initial_followup_messages = st.session_state.thread_with_tavily_context
                        else:
                            st.session_state.initial_response_thread.append({"role": "user", "content": initial_followup_question})
                            initial_followup_messages = st.session_state.initial_response_thread
                        with st.chat_message("user"):
                            st.markdown(initial_followup_question)
                        with st.chat_message("assistant"):
                            client = OpenAI()
                            try:
                                stream = client.chat.completions.create(
                                    model="gpt-4o",
                                    messages=[{"role": m["role"], "content": m["content"]} for m in initial_followup_messages],
                                    stream=True,
                                )
                                response = st.write_stream(stream)
                            except Exception as e:
                                st.error(f"Error during OpenAI call: {e}")
                                return
                            st.session_state.initial_response_thread.append({"role": "assistant", "content": response})
                    with st.expander("View full follow-up thread"):
                        for message in st.session_state.initial_response_thread:
                            if message['role'] != 'system':
                                emoji = role_emojis.get(message['role'], "❓")
                                st.write(f"{emoji} {message['role'].capitalize()}: {message['content']}")
                    if st.session_state.initial_response_thread:
                        if st.checkbox("Download Followup Conversation"):
                            full_followup_conversation = ""
                            for message in st.session_state.initial_response_thread:
                                if message['role'] != 'system':
                                    emoji = role_emojis.get(message['role'], "❓")
                                    full_followup_conversation += f"{emoji} {message['role'].capitalize()}: {message['content']}\n\n"
                            html = markdown2.markdown(full_followup_conversation, extras=["tables"])
                            st.download_button('Download Followup Conversation', html, 'followup_conversation.html', 'text/html')
                if not initial_followup:
                    if st.button("Ask 3 AI Expert Personas for Opinions"):
                        prelim_response = st.session_state.rag_response + st.session_state.source_chunks
                        try:
                            completion = create_chat_completion(
                                messages=find_experts_messages,
                                model=experts_model,
                                temperature=0.3,
                                response_format="json_object"
                            )
                        except Exception as e:
                            st.error(f"Error during OpenAI call: {e}")
                            return
                        json_output = completion.choices[0].message.content
                        experts, domains, expert_questions = extract_expert_info(json_output)
                        st.session_state.experts = experts
                        updated_expert1_system_prompt = expert1_system_prompt.format(expert=experts[0], domain=domains[0])
                        updated_expert2_system_prompt = expert2_system_prompt.format(expert=experts[1], domain=domains[1])
                        updated_expert3_system_prompt = expert3_system_prompt.format(expert=experts[2], domain=domains[2])
                        updated_question1 = expert_questions[0]
                        updated_question2 = expert_questions[1]
                        updated_question3 = expert_questions[2]
                        prelim_response = st.session_state.rag_response + st.session_state.source_chunks
                        expert1_messages = [
                            {'role': 'system', 'content': updated_expert1_system_prompt},
                            {'role': 'user', 'content': updated_question1 + "Here's what I already found online: " + prelim_response}
                        ]
                        st.session_state.messages1 = expert1_messages
                        expert2_messages = [
                            {'role': 'system', 'content': updated_expert2_system_prompt},
                            {'role': 'user', 'content': updated_question2 + "Here's what I already found online: " + prelim_response}
                        ]
                        st.session_state.messages2 = expert2_messages
                        expert3_messages = [
                            {'role': 'system', 'content': updated_expert3_system_prompt},
                            {'role': 'user', 'content': updated_question3 + "Here's what I already found online: " + prelim_response}
                        ]
                        st.session_state.messages3 = expert3_messages
                        with st.spinner('Waiting for experts to respond...'):
                            st.session_state.expert_answers = asyncio.run(
                                get_responses([expert1_messages, expert2_messages, expert3_messages])
                            )
        with col1:
            if not first_view:
                try:
                    if st.session_state.pubmed_search_terms and st.session_state.articles:
                        with st.expander("View PubMed Results Added to Knowledge Base"):
                            pubmed_link = "https://pubmed.ncbi.nlm.nih.gov/?term=" + st.session_state.pubmed_search_terms
                            st.page_link(pubmed_link, label="Click here to view in PubMed", icon="📚")
                            with st.popover("PubMed Search Terms"):
                                st.write(f'**Search Strategy:** {st.session_state.pubmed_search_terms}')
                            for article in st.session_state.articles:
                                st.markdown(f"### [{article['title']}]({article['link']})")
                                st.write(f"Year: {article['year']}")
                                st.write(article['abstract'] if article['abstract'] else "No abstract available")
                except:
                    st.write("No Relevant PubMed articles to display - if topic works, API may be down!")
                    pubmed_link = "https://pubmed.ncbi.nlm.nih.gov/?term=" + st.session_state.pubmed_search_terms
                    st.page_link(pubmed_link, label="Click here to try in PubMed", icon="📚")
                    with st.popover("PubMed Search Terms"):
                        st.write(f'**Search Strategy:** {st.session_state.pubmed_search_terms}')
                with st.expander("View Internet Results Added to Knowledge Base"):
                    for snippet in st.session_state.snippets:
                        st.markdown(snippet.replace('<END OF SITE>', ''))
                if st.session_state.rag_response:
                    st.info("Initial Answer")
                    container1 = st.container()
                    with container1:
                        st.markdown(st.session_state.rag_response)
                        if st.button("Create Word File"):
                            doc = markdown_to_word(st.session_state.rag_response)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File",
                                data=buffer,
                                file_name="prelim_response.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        with st.expander("View Source Excerpts"):
                            st.markdown(st.session_state.source_chunks)
        with col2:
            if st.session_state.expert_answers:
                container2 = st.container()
                with container2:
                    st.info("AI Expert Persona Responses")
                    with st.expander(f'AI {st.session_state.experts[0]} Perspective'):
                        expert_0 = st.session_state.expert_answers[0]['choices'][0]['message']['content']
                        st.write(expert_0)
                        st.session_state.messages1.append({"role": "assistant", "content": expert_0})
                        if st.button("Create Word File for AI Expert 1"):
                            doc = markdown_to_word(expert_0)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File for AI Expert 1",
                                data=buffer,
                                file_name="AI_expert_1.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    with st.expander(f'AI {st.session_state.experts[1]} Perspective'):
                        expert_1 = st.session_state.expert_answers[1]['choices'][0]['message']['content']
                        st.write(expert_1)
                        st.session_state.messages2.append({"role": "assistant", "content": expert_1})
                        if st.button("Create Word File for AI Expert 2"):
                            doc = markdown_to_word(expert_1)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File for AI Expert 2",
                                data=buffer,
                                file_name="AI_expert_2.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    with st.expander(f'AI {st.session_state.experts[2]} Perspective'):
                        expert_2 = st.session_state.expert_answers[2]['choices'][0]['message']['content']
                        st.write(expert_2)
                        st.session_state.messages3.append({"role": "assistant", "content": expert_2})
                        if st.button("Create Word File for AI Expert 3"):
                            doc = markdown_to_word(expert_2)
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="Download Word File for AI Expert 3",
                                data=buffer,
                                file_name="AI_expert_3.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

#########################################
# Main Program Entry Point
#########################################
if __name__ == '__main__':
    main()
